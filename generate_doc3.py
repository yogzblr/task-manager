#!/usr/bin/env python3
"""Generate Document 3: Implementation Guide"""

from docx import Document
from docx.shared import Pt

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph(code_text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

print("Generating Document 3: Implementation Guide...")
doc3 = Document()
doc3.add_heading('Multi-Tenant VM Manager - Implementation Guide', 0)

# Executive Summary
add_heading(doc3, '1. Executive Summary', 1)
add_paragraph(doc3, 'This comprehensive implementation guide provides step-by-step instructions for building the Multi-Tenant VM Management system. It covers development environment setup, detailed component implementation, testing strategies, and deployment procedures.')
add_paragraph(doc3, '')
add_paragraph(doc3, 'Target Audience:', bold=True)
add_paragraph(doc3, '• Backend developers (Go experience required)')
add_paragraph(doc3, '• DevOps engineers')
add_paragraph(doc3, '• System architects')
add_paragraph(doc3, '')
add_paragraph(doc3, 'Implementation Timeline:', bold=True)
add_paragraph(doc3, '• Phase 1: Core agent (4 weeks)')
add_paragraph(doc3, '• Phase 2: Control plane (4 weeks)')
add_paragraph(doc3, '• Phase 3: MCP integration (2 weeks)')
add_paragraph(doc3, '• Phase 4: Testing & docs (2 weeks)')
add_paragraph(doc3, '• Total: 12 weeks')

# Development Environment Setup
add_heading(doc3, '2. Development Environment Setup', 1)

add_heading(doc3, '2.1 Required Tools', 2)
add_paragraph(doc3, 'Install the following tools:', bold=True)
add_paragraph(doc3, '')
add_code_block(doc3, '''# Go 1.21 or higher
wget https://go.dev/dl/go1.21.0.linux-amd64.tar.gz
sudo tar -C /usr/local -xzf go1.21.0.linux-amd64.tar.gz
export PATH=$PATH:/usr/local/go/bin

# Docker
curl -fsSL https://get.docker.com -o get-docker.sh
sudo sh get-docker.sh

# kubectl
curl -LO "https://dl.k8s.io/release/$(curl -L -s https://dl.k8s.io/release/stable.txt)/bin/linux/amd64/kubectl"
sudo install -o root -g root -m 0755 kubectl /usr/local/bin/kubectl

# kind (local Kubernetes)
curl -Lo ./kind https://kind.sigs.k8s.io/dl/v0.20.0/kind-linux-amd64
chmod +x ./kind
sudo mv ./kind /usr/local/bin/kind

# MySQL client
sudo apt-get install -y mysql-client

# Development tools
go install github.com/golangci/golangci-lint/cmd/golangci-lint@latest
go install github.com/swaggo/swag/cmd/swag@latest
''')

add_heading(doc3, '2.2 Local Development Cluster', 2)
add_code_block(doc3, '''# Create kind cluster
cat <<EOF | kind create cluster --config=-
kind: Cluster
apiVersion: kind.x-k8s.io/v1alpha4
name: vm-manager-dev
nodes:
- role: control-plane
- role: worker
- role: worker
EOF

# Verify cluster
kubectl cluster-info
kubectl get nodes
''')

add_heading(doc3, '2.3 Go Module Initialization', 2)
add_code_block(doc3, '''# Create workspace
mkdir -p ~/vm-manager/{vm-agent,control-plane,deploy}
cd ~/vm-manager

# Initialize agent module
cd vm-agent
go mod init github.com/yourorg/vm-agent
go mod tidy

# Initialize control plane module
cd ../control-plane
go mod init github.com/yourorg/control-plane
go mod tidy
''')

# Agent Implementation
add_heading(doc3, '3. Agent Implementation', 1)

add_heading(doc3, '3.1 Project Structure', 2)
add_code_block(doc3, '''vm-agent/
├── cmd/
│   └── agent/
│       └── main.go                 # Entry point
├── pkg/
│   ├── agent/                      # Agent manager
│   │   ├── manager.go
│   │   ├── coordinator.go
│   │   └── context.go
│   ├── piko/                       # Piko client
│   │   ├── client.go
│   │   ├── tunnel.go
│   │   └── reconnect.go
│   ├── webhook/                    # Webhook server
│   │   ├── server.go
│   │   ├── handlers.go
│   │   └── auth.go
│   ├── probe/                      # Probe executor
│   │   ├── executor.go
│   │   ├── workflow.go
│   │   └── reporter.go
│   ├── config/                     # Configuration
│   │   ├── loader.go
│   │   ├── priority.go
│   │   └── validator.go
│   ├── health/                     # Health monitoring
│   │   ├── monitor.go
│   │   ├── checks.go
│   │   └── reporter.go
│   └── lifecycle/                  # Lifecycle management
│       ├── install.go
│       ├── configure.go
│       ├── repair.go
│       ├── upgrade.go
│       ├── uninstall.go
│       ├── service_linux.go
│       └── service_windows.go
├── internal/
│   └── version/
│       └── version.go              # Version information
├── Makefile                        # Build automation
├── go.mod
├── go.sum
└── README.md
''')

add_heading(doc3, '3.2 Agent Manager Implementation', 2)
add_paragraph(doc3, 'File: pkg/agent/manager.go', bold=True)
add_code_block(doc3, '''package agent

import (
    "context"
    "fmt"
    "os"
    "os/signal"
    "sync"
    "syscall"

    "github.com/spf13/cobra"
    "github.com/yourorg/vm-agent/pkg/config"
    "github.com/yourorg/vm-agent/pkg/health"
    "github.com/yourorg/vm-agent/pkg/lifecycle"
    "github.com/yourorg/vm-agent/pkg/piko"
    "github.com/yourorg/vm-agent/pkg/probe"
    "github.com/yourorg/vm-agent/pkg/webhook"
)

type Manager struct {
    cfg         *config.Config
    pikoClient  *piko.Client
    webhookSrv  *webhook.Server
    probeExec   *probe.Executor
    healthMon   *health.Monitor
    ctx         context.Context
    cancel      context.CancelFunc
    wg          sync.WaitGroup
}

func NewManager(cfg *config.Config) (*Manager, error) {
    ctx, cancel := context.WithCancel(context.Background())

    m := &Manager{
        cfg:    cfg,
        ctx:    ctx,
        cancel: cancel,
    }

    // Initialize components
    var err error

    m.pikoClient, err = piko.NewClient(cfg.Piko)
    if err != nil {
        return nil, fmt.Errorf("failed to create piko client: %w", err)
    }

    m.probeExec = probe.NewExecutor(cfg.Probe)

    m.webhookSrv, err = webhook.NewServer(cfg.Webhook, m.probeExec)
    if err != nil {
        return nil, fmt.Errorf("failed to create webhook server: %w", err)
    }

    m.healthMon = health.NewMonitor(cfg.Health, m)

    return m, nil
}

func (m *Manager) Run() error {
    // Start Piko tunnel
    m.wg.Add(1)
    go func() {
        defer m.wg.Done()
        if err := m.pikoClient.Connect(m.ctx); err != nil {
            fmt.Printf("Piko client error: %v\\n", err)
        }
    }()

    // Start webhook server
    m.wg.Add(1)
    go func() {
        defer m.wg.Done()
        if err := m.webhookSrv.Start(m.ctx); err != nil {
            fmt.Printf("Webhook server error: %v\\n", err)
        }
    }()

    // Start health monitor
    m.wg.Add(1)
    go func() {
        defer m.wg.Done()
        m.healthMon.Run(m.ctx)
    }()

    // Wait for shutdown signal
    sigCh := make(chan os.Signal, 1)
    signal.Notify(sigCh, os.Interrupt, syscall.SIGTERM)

    select {
    case <-sigCh:
        fmt.Println("Shutdown signal received")
    case <-m.ctx.Done():
        fmt.Println("Context cancelled")
    }

    // Graceful shutdown
    m.cancel()
    m.wg.Wait()

    return nil
}

// CLI commands
func (m *Manager) RootCmd() *cobra.Command {
    rootCmd := &cobra.Command{
        Use:   "vm-agent",
        Short: "Multi-tenant VM management agent",
    }

    rootCmd.AddCommand(
        m.InstallCmd(),
        m.ConfigureCmd(),
        m.RepairCmd(),
        m.UpgradeCmd(),
        m.UninstallCmd(),
        m.StatusCmd(),
        m.RunCmd(),
    )

    return rootCmd
}

func (m *Manager) RunCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "run",
        Short: "Run the agent service",
        RunE: func(cmd *cobra.Command, args []string) error {
            return m.Run()
        },
    }
}

func (m *Manager) InstallCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "install",
        Short: "Install agent as system service",
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Install(m.cfg)
        },
    }
}

func (m *Manager) ConfigureCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "configure",
        Short: "Update agent configuration",
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Configure(m.cfg)
        },
    }
}

func (m *Manager) RepairCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "repair",
        Short: "Run diagnostics and repair",
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Repair(m.cfg)
        },
    }
}

func (m *Manager) UpgradeCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "upgrade <version>",
        Short: "Upgrade agent to new version",
        Args:  cobra.ExactArgs(1),
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Upgrade(m.cfg, args[0])
        },
    }
}

func (m *Manager) UninstallCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "uninstall",
        Short: "Uninstall agent and remove service",
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Uninstall(m.cfg)
        },
    }
}

func (m *Manager) StatusCmd() *cobra.Command {
    return &cobra.Command{
        Use:   "status",
        Short: "Show agent status and health",
        RunE: func(cmd *cobra.Command, args []string) error {
            return lifecycle.Status(m.cfg)
        },
    }
}

// Health check interface
func (m *Manager) HealthCheck() health.Status {
    status := health.Status{
        Healthy: true,
        Components: make(map[string]health.ComponentStatus),
    }

    // Check Piko client
    if m.pikoClient.IsConnected() {
        status.Components["piko"] = health.ComponentStatus{
            Healthy: true,
            Message: "Connected",
        }
    } else {
        status.Healthy = false
        status.Components["piko"] = health.ComponentStatus{
            Healthy: false,
            Message: "Disconnected",
        }
    }

    // Check webhook server
    if m.webhookSrv.IsHealthy() {
        status.Components["webhook"] = health.ComponentStatus{
            Healthy: true,
            Message: "Running",
        }
    } else {
        status.Healthy = false
        status.Components["webhook"] = health.ComponentStatus{
            Healthy: false,
            Message: "Not responding",
        }
    }

    return status
}
''')

add_heading(doc3, '3.3 Configuration Management', 2)
add_paragraph(doc3, 'File: pkg/config/loader.go', bold=True)
add_code_block(doc3, '''package config

import (
    "fmt"
    "os"
    "time"

    "github.com/spf13/viper"
)

type Config struct {
    Agent    AgentConfig
    Piko     PikoConfig
    Webhook  WebhookConfig
    Probe    ProbeConfig
    Health   HealthConfig
}

type AgentConfig struct {
    ID              string
    TenantID        string
    InstallKey      string
    Token           string
    ControlPlaneURL string
}

type PikoConfig struct {
    ServerURL string
    Endpoint  string
    Token     string
    Reconnect ReconnectConfig
}

type ReconnectConfig struct {
    InitialDelay time.Duration
    MaxDelay     time.Duration
    Multiplier   float64
}

type WebhookConfig struct {
    ListenAddr string
    Port       int
    TLSEnabled bool
    CertFile   string
    KeyFile    string
}

type ProbeConfig struct {
    WorkDir         string
    DefaultTimeout  time.Duration
    MaxConcurrent   int
}

type HealthConfig struct {
    CheckInterval   time.Duration
    ReportInterval  time.Duration
    ReportURL       string
}

func Load() (*Config, error) {
    v := viper.New()

    // Set defaults
    v.SetDefault("agent.id", "")
    v.SetDefault("piko.reconnect.initial_delay", "1s")
    v.SetDefault("piko.reconnect.max_delay", "60s")
    v.SetDefault("piko.reconnect.multiplier", 2.0)
    v.SetDefault("webhook.listen_addr", "0.0.0.0")
    v.SetDefault("webhook.port", 9999)
    v.SetDefault("webhook.tls_enabled", false)
    v.SetDefault("probe.work_dir", "/var/lib/vm-agent/work")
    v.SetDefault("probe.default_timeout", "300s")
    v.SetDefault("probe.max_concurrent", 5)
    v.SetDefault("health.check_interval", "30s")
    v.SetDefault("health.report_interval", "300s")

    // Environment variables (highest priority)
    v.SetEnvPrefix("VM_AGENT")
    v.AutomaticEnv()

    // Configuration file
    v.SetConfigName("config")
    v.SetConfigType("yaml")
    v.AddConfigPath("/etc/vm-agent/")
    v.AddConfigPath("$HOME/.vm-agent/")
    v.AddConfigPath(".")

    if err := v.ReadInConfig(); err != nil {
        if _, ok := err.(viper.ConfigFileNotFoundError); !ok {
            return nil, fmt.Errorf("failed to read config: %w", err)
        }
        // Config file not found is OK, will use env vars and defaults
    }

    var cfg Config
    if err := v.Unmarshal(&cfg); err != nil {
        return nil, fmt.Errorf("failed to unmarshal config: %w", err)
    }

    // Validate configuration
    if err := Validate(&cfg); err != nil {
        return nil, fmt.Errorf("invalid configuration: %w", err)
    }

    return &cfg, nil
}

func Validate(cfg *Config) error {
    if cfg.Agent.ID == "" {
        return fmt.Errorf("agent.id is required")
    }
    if cfg.Agent.TenantID == "" {
        return fmt.Errorf("agent.tenant_id is required")
    }
    if cfg.Piko.ServerURL == "" {
        return fmt.Errorf("piko.server_url is required")
    }
    if cfg.Piko.Token == "" && cfg.Agent.Token == "" {
        return fmt.Errorf("either piko.token or agent.token is required")
    }
    return nil
}
''')

add_heading(doc3, '3.4 Lifecycle Commands - Linux systemd', 2)
add_paragraph(doc3, 'File: pkg/lifecycle/service_linux.go', bold=True)
add_code_block(doc3, '''package lifecycle

import (
    "fmt"
    "os"
    "os/exec"
    "path/filepath"

    "github.com/yourorg/vm-agent/pkg/config"
)

const systemdServiceTemplate = `[Unit]
Description=VM Agent
Documentation=https://docs.example.com/vm-agent
After=network-online.target
Wants=network-online.target

[Service]
Type=simple
User=vm-agent
Group=vm-agent
ExecStart=%s run
ExecReload=/bin/kill -HUP $MAINPID
Restart=always
RestartSec=10
LimitNOFILE=65536

# Environment
Environment="VM_AGENT_ID=%s"
Environment="VM_AGENT_TENANT_ID=%s"
Environment="VM_AGENT_TOKEN=%s"
Environment="VM_AGENT_PIKO_SERVER_URL=%s"

# Security
NoNewPrivileges=true
PrivateTmp=true
ProtectSystem=strict
ProtectHome=true
ReadWritePaths=/var/lib/vm-agent /var/log/vm-agent

[Install]
WantedBy=multi-user.target
`

func Install(cfg *config.Config) error {
    // Check if running as root
    if os.Geteuid() != 0 {
        return fmt.Errorf("must run as root")
    }

    // Copy binary to /usr/local/bin
    exePath, err := os.Executable()
    if err != nil {
        return fmt.Errorf("failed to get executable path: %w", err)
    }

    targetPath := "/usr/local/bin/vm-agent"
    if err := copyFile(exePath, targetPath); err != nil {
        return fmt.Errorf("failed to copy binary: %w", err)
    }

    if err := os.Chmod(targetPath, 0755); err != nil {
        return fmt.Errorf("failed to chmod binary: %w", err)
    }

    // Create vm-agent user and group
    if err := createUser(); err != nil {
        return fmt.Errorf("failed to create user: %w", err)
    }

    // Create directories
    dirs := []string{
        "/etc/vm-agent",
        "/var/lib/vm-agent",
        "/var/lib/vm-agent/work",
        "/var/log/vm-agent",
    }

    for _, dir := range dirs {
        if err := os.MkdirAll(dir, 0755); err != nil {
            return fmt.Errorf("failed to create dir %s: %w", dir, err)
        }
        if err := os.Chown(dir, lookupUID("vm-agent"), lookupGID("vm-agent")); err != nil {
            return fmt.Errorf("failed to chown dir %s: %w", dir, err)
        }
    }

    // Write systemd service file
    serviceContent := fmt.Sprintf(systemdServiceTemplate,
        targetPath,
        cfg.Agent.ID,
        cfg.Agent.TenantID,
        cfg.Agent.Token,
        cfg.Piko.ServerURL,
    )

    servicePath := "/etc/systemd/system/vm-agent.service"
    if err := os.WriteFile(servicePath, []byte(serviceContent), 0644); err != nil {
        return fmt.Errorf("failed to write service file: %w", err)
    }

    // Reload systemd
    if err := exec.Command("systemctl", "daemon-reload").Run(); err != nil {
        return fmt.Errorf("failed to reload systemd: %w", err)
    }

    // Enable service
    if err := exec.Command("systemctl", "enable", "vm-agent").Run(); err != nil {
        return fmt.Errorf("failed to enable service: %w", err)
    }

    // Start service
    if err := exec.Command("systemctl", "start", "vm-agent").Run(); err != nil {
        return fmt.Errorf("failed to start service: %w", err)
    }

    fmt.Println("Agent installed and started successfully")
    return nil
}

func createUser() error {
    // Check if user exists
    if _, err := exec.Command("id", "vm-agent").Output(); err == nil {
        return nil // User already exists
    }

    // Create system user
    cmd := exec.Command("useradd",
        "--system",
        "--no-create-home",
        "--shell", "/bin/false",
        "vm-agent",
    )

    return cmd.Run()
}

func copyFile(src, dst string) error {
    input, err := os.ReadFile(src)
    if err != nil {
        return err
    }
    return os.WriteFile(dst, input, 0755)
}

func lookupUID(username string) int {
    // Simplified - use proper user lookup in production
    return 999
}

func lookupGID(groupname string) int {
    // Simplified - use proper group lookup in production
    return 999
}

func Uninstall(cfg *config.Config) error {
    // Check if running as root
    if os.Geteuid() != 0 {
        return fmt.Errorf("must run as root")
    }

    // Stop service
    exec.Command("systemctl", "stop", "vm-agent").Run()

    // Disable service
    exec.Command("systemctl", "disable", "vm-agent").Run()

    // Remove service file
    os.Remove("/etc/systemd/system/vm-agent.service")

    // Reload systemd
    exec.Command("systemctl", "daemon-reload").Run()

    // Remove binary
    os.Remove("/usr/local/bin/vm-agent")

    // Optionally remove data directories
    fmt.Println("Agent uninstalled successfully")
    fmt.Println("Data directories preserved: /var/lib/vm-agent, /var/log/vm-agent")
    fmt.Println("Run 'rm -rf /var/lib/vm-agent /var/log/vm-agent' to remove all data")

    return nil
}
''')

add_heading(doc3, '3.5 Self-Upgrade Implementation', 2)
add_paragraph(doc3, 'File: pkg/lifecycle/upgrade.go', bold=True)
add_code_block(doc3, '''package lifecycle

import (
    "crypto/sha256"
    "fmt"
    "io"
    "net/http"
    "os"
    "os/exec"
    "path/filepath"
    "time"

    "github.com/yourorg/vm-agent/pkg/config"
)

type UpgradeManager struct {
    cfg            *config.Config
    downloadURL    string
    targetVersion  string
    workDir        string
}

func Upgrade(cfg *config.Config, targetVersion string) error {
    um := &UpgradeManager{
        cfg:           cfg,
        targetVersion: targetVersion,
        workDir:       "/var/lib/vm-agent/upgrade",
    }

    fmt.Printf("Upgrading to version %s...\\n", targetVersion)

    // Create work directory
    if err := os.MkdirAll(um.workDir, 0755); err != nil {
        return fmt.Errorf("failed to create work dir: %w", err)
    }

    // Step 1: Download new version
    newBinaryPath, err := um.download()
    if err != nil {
        return fmt.Errorf("download failed: %w", err)
    }

    // Step 2: Verify checksum
    if err := um.verify(newBinaryPath); err != nil {
        return fmt.Errorf("verification failed: %w", err)
    }

    // Step 3: Backup current version
    backupPath, err := um.backup()
    if err != nil {
        return fmt.Errorf("backup failed: %w", err)
    }

    // Step 4: Replace binary
    if err := um.replace(newBinaryPath); err != nil {
        return fmt.Errorf("replace failed: %w", err)
    }

    // Step 5: Restart service
    if err := um.restart(); err != nil {
        fmt.Println("Restart failed, attempting rollback...")
        if rollbackErr := um.rollback(backupPath); rollbackErr != nil {
            return fmt.Errorf("restart failed and rollback failed: %w", rollbackErr)
        }
        return fmt.Errorf("upgrade failed, rolled back: %w", err)
    }

    // Step 6: Verify new version is running
    time.Sleep(5 * time.Second)
    if err := um.verifyRunning(); err != nil {
        fmt.Println("Verification failed, attempting rollback...")
        if rollbackErr := um.rollback(backupPath); rollbackErr != nil {
            return fmt.Errorf("verification failed and rollback failed: %w", rollbackErr)
        }
        return fmt.Errorf("upgrade verification failed, rolled back: %w", err)
    }

    // Cleanup
    os.RemoveAll(um.workDir)

    fmt.Printf("Upgrade to version %s completed successfully\\n", targetVersion)
    return nil
}

func (um *UpgradeManager) download() (string, error) {
    // Construct download URL
    um.downloadURL = fmt.Sprintf(
        "%s/downloads/vm-agent/%s/vm-agent-%s-%s",
        um.cfg.Agent.ControlPlaneURL,
        um.targetVersion,
        um.targetVersion,
        getOSArch(),
    )

    fmt.Printf("Downloading from %s...\\n", um.downloadURL)

    resp, err := http.Get(um.downloadURL)
    if err != nil {
        return "", fmt.Errorf("download request failed: %w", err)
    }
    defer resp.Body.Close()

    if resp.StatusCode != 200 {
        return "", fmt.Errorf("download failed with status: %d", resp.StatusCode)
    }

    // Save to file
    outputPath := filepath.Join(um.workDir, "vm-agent-new")
    out, err := os.Create(outputPath)
    if err != nil {
        return "", fmt.Errorf("failed to create output file: %w", err)
    }
    defer out.Close()

    if _, err := io.Copy(out, resp.Body); err != nil {
        return "", fmt.Errorf("failed to write file: %w", err)
    }

    // Make executable
    if err := os.Chmod(outputPath, 0755); err != nil {
        return "", fmt.Errorf("failed to chmod: %w", err)
    }

    fmt.Println("Download completed")
    return outputPath, nil
}

func (um *UpgradeManager) verify(binaryPath string) error {
    // Download checksum
    checksumURL := um.downloadURL + ".sha256"
    resp, err := http.Get(checksumURL)
    if err != nil {
        return fmt.Errorf("failed to download checksum: %w", err)
    }
    defer resp.Body.Close()

    expectedChecksum, err := io.ReadAll(resp.Body)
    if err != nil {
        return fmt.Errorf("failed to read checksum: %w", err)
    }

    // Calculate actual checksum
    f, err := os.Open(binaryPath)
    if err != nil {
        return fmt.Errorf("failed to open binary: %w", err)
    }
    defer f.Close()

    h := sha256.New()
    if _, err := io.Copy(h, f); err != nil {
        return fmt.Errorf("failed to calculate checksum: %w", err)
    }

    actualChecksum := fmt.Sprintf("%x", h.Sum(nil))

    if actualChecksum != string(expectedChecksum) {
        return fmt.Errorf("checksum mismatch: expected %s, got %s",
            string(expectedChecksum), actualChecksum)
    }

    fmt.Println("Checksum verified")
    return nil
}

func (um *UpgradeManager) backup() (string, error) {
    currentPath := "/usr/local/bin/vm-agent"
    backupPath := filepath.Join(um.workDir, "vm-agent-backup")

    input, err := os.ReadFile(currentPath)
    if err != nil {
        return "", fmt.Errorf("failed to read current binary: %w", err)
    }

    if err := os.WriteFile(backupPath, input, 0755); err != nil {
        return "", fmt.Errorf("failed to write backup: %w", err)
    }

    fmt.Println("Backup created")
    return backupPath, nil
}

func (um *UpgradeManager) replace(newBinaryPath string) error {
    targetPath := "/usr/local/bin/vm-agent"

    // Atomic replace using rename
    if err := os.Rename(newBinaryPath, targetPath); err != nil {
        return fmt.Errorf("failed to replace binary: %w", err)
    }

    fmt.Println("Binary replaced")
    return nil
}

func (um *UpgradeManager) restart() error {
    fmt.Println("Restarting service...")
    return exec.Command("systemctl", "restart", "vm-agent").Run()
}

func (um *UpgradeManager) verifyRunning() error {
    // Check service status
    out, err := exec.Command("systemctl", "is-active", "vm-agent").Output()
    if err != nil {
        return fmt.Errorf("service not active: %w", err)
    }

    if string(out) != "active\\n" {
        return fmt.Errorf("service not active: %s", string(out))
    }

    // Verify version
    out, err = exec.Command("/usr/local/bin/vm-agent", "status").Output()
    if err != nil {
        return fmt.Errorf("failed to get version: %w", err)
    }

    // Parse version from output
    // ... implementation details ...

    fmt.Println("New version verified")
    return nil
}

func (um *UpgradeManager) rollback(backupPath string) error {
    fmt.Println("Rolling back to previous version...")

    targetPath := "/usr/local/bin/vm-agent"
    if err := os.Rename(backupPath, targetPath); err != nil {
        return fmt.Errorf("failed to restore backup: %w", err)
    }

    if err := exec.Command("systemctl", "restart", "vm-agent").Run(); err != nil {
        return fmt.Errorf("failed to restart after rollback: %w", err)
    }

    fmt.Println("Rollback completed")
    return nil
}

func getOSArch() string {
    // Return OS-arch string like "linux-amd64", "windows-amd64", etc.
    // Implementation details...
    return "linux-amd64"
}
''')

# Save Document 3
doc3.save('doc3_implementation.docx')
print("Document 3 completed: doc3_implementation.docx")

#!/usr/bin/env python3
"""Generate Document 4: Claude Code Specification"""

from docx import Document
from docx.shared import Pt

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph(code_text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

print("Generating Document 4: Claude Code Specification...")
doc4 = Document()
doc4.add_heading('Multi-Tenant VM Manager - Claude Code Specification', 0)

# Project Overview
add_heading(doc4, '1. Project Overview', 1)
add_paragraph(doc4, 'This document provides complete specifications for Claude Code to generate the Multi-Tenant VM Management system. It includes technology stack, directory structure, dependencies, build instructions, and detailed code generation tasks.')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Purpose:', bold=True)
add_paragraph(doc4, 'Enable Claude Code to autonomously generate production-ready code for all system components including the unified agent, control plane API, MCP server, database schema, and Kubernetes deployment manifests.')

# Technology Stack
add_heading(doc4, '2. Technology Stack', 1)

add_paragraph(doc4, 'Programming Languages:', bold=True)
add_paragraph(doc4, '• Go 1.21+ (primary language for all backend components)')
add_paragraph(doc4, '• Python 3.10+ (for MCP server)')
add_paragraph(doc4, '• SQL (MySQL 8.0+ or PostgreSQL 13+)')
add_paragraph(doc4, '• YAML (configuration and workflows)')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Core Dependencies:', bold=True)
add_paragraph(doc4, '• Piko v0.8+ with multi-tenant auth and rebalancing PRs')
add_paragraph(doc4, '• adnanh/webhook (embedded in agent)')
add_paragraph(doc4, '• linyows/probe (embedded in agent)')
add_paragraph(doc4, '• Quickwit latest (audit logging)')
add_paragraph(doc4, '• Kubernetes 1.27+ (orchestration)')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Go Libraries:', bold=True)
add_code_block(doc4, '''# Core
github.com/spf13/cobra           # CLI framework
github.com/spf13/viper           # Configuration management
github.com/golang-jwt/jwt/v5     # JWT authentication

# HTTP/Web
github.com/gin-gonic/gin         # REST API framework
github.com/gorilla/websocket     # WebSocket support

# Database
github.com/go-sql-driver/mysql   # MySQL driver
gorm.io/gorm                     # ORM
gorm.io/driver/mysql             # GORM MySQL driver

# Kubernetes
k8s.io/client-go                 # Kubernetes client
k8s.io/apimachinery              # Kubernetes API machinery

# Observability
github.com/prometheus/client_golang  # Prometheus metrics
go.uber.org/zap                      # Structured logging

# Testing
github.com/stretchr/testify      # Testing framework
github.com/golang/mock           # Mocking framework
''')

# Directory Structure
add_heading(doc4, '3. Complete Directory Structure', 1)

add_code_block(doc4, '''vm-manager/
├── vm-agent/                           # Agent component
│   ├── cmd/
│   │   └── agent/
│   │       └── main.go
│   ├── pkg/
│   │   ├── agent/
│   │   │   ├── manager.go              # Main agent manager
│   │   │   ├── coordinator.go          # Component coordination
│   │   │   └── context.go              # Shared context
│   │   ├── piko/
│   │   │   ├── client.go               # Piko client wrapper
│   │   │   ├── tunnel.go               # Tunnel management
│   │   │   └── reconnect.go            # Reconnection logic
│   │   ├── webhook/
│   │   │   ├── server.go               # HTTP server
│   │   │   ├── handlers.go             # Request handlers
│   │   │   └── auth.go                 # Authentication
│   │   ├── probe/
│   │   │   ├── executor.go             # Workflow executor
│   │   │   ├── workflow.go             # Workflow parsing
│   │   │   └── reporter.go             # Result reporting
│   │   ├── config/
│   │   │   ├── loader.go               # Config loading
│   │   │   ├── priority.go             # Priority resolution
│   │   │   └── validator.go            # Validation
│   │   ├── health/
│   │   │   ├── monitor.go              # Health monitoring
│   │   │   ├── checks.go               # Health checks
│   │   │   └── reporter.go             # Health reporting
│   │   └── lifecycle/
│   │       ├── install.go              # Installation
│   │       ├── configure.go            # Configuration
│   │       ├── repair.go               # Self-repair
│   │       ├── upgrade.go              # Self-upgrade
│   │       ├── uninstall.go            # Uninstallation
│   │       ├── service_linux.go        # Linux systemd
│   │       └── service_windows.go      # Windows service
│   ├── internal/
│   │   └── version/
│   │       └── version.go              # Version info
│   ├── Dockerfile
│   ├── Makefile
│   ├── go.mod
│   ├── go.sum
│   └── README.md
│
├── control-plane/                      # Control plane API
│   ├── cmd/
│   │   └── server/
│   │       └── main.go
│   ├── pkg/
│   │   ├── api/
│   │   │   ├── server.go               # HTTP server
│   │   │   ├── middleware.go           # Middleware
│   │   │   ├── router.go               # Route definitions
│   │   │   └── handlers/
│   │   │       ├── tenants.go          # Tenant handlers
│   │   │       ├── agents.go           # Agent handlers
│   │   │       ├── workflows.go        # Workflow handlers
│   │   │       ├── campaigns.go        # Campaign handlers
│   │   │       └── audit.go            # Audit handlers
│   │   ├── tenant/
│   │   │   ├── manager.go              # Tenant management
│   │   │   └── isolation.go            # Tenant isolation
│   │   ├── agent/
│   │   │   ├── registry.go             # Agent registry
│   │   │   ├── registration.go         # Registration flow
│   │   │   └── keys.go                 # Installation keys
│   │   ├── workflow/
│   │   │   ├── manager.go              # Workflow manager
│   │   │   ├── validator.go            # Workflow validation
│   │   │   └── executor.go             # Workflow execution
│   │   ├── campaign/
│   │   │   ├── manager.go              # Campaign orchestration
│   │   │   ├── phases.go               # Phased rollout
│   │   │   ├── tracker.go              # Progress tracking
│   │   │   └── rollback.go             # Auto-rollback
│   │   ├── mcp/
│   │   │   ├── server.go               # MCP server
│   │   │   ├── tools.go                # Tool definitions
│   │   │   ├── generator.go            # Workflow generation
│   │   │   └── validator.go            # Validation
│   │   ├── audit/
│   │   │   ├── quickwit.go             # Quickwit client
│   │   │   ├── schema.go               # Log schema
│   │   │   └── logger.go               # Audit logger
│   │   ├── db/
│   │   │   ├── connection.go           # DB connection
│   │   │   ├── migration.go            # Migration runner
│   │   │   └── models/
│   │   │       ├── tenant.go           # Tenant model
│   │   │       ├── agent.go            # Agent model
│   │   │       ├── key.go              # Key model
│   │   │       ├── workflow.go         # Workflow model
│   │   │       └── campaign.go         # Campaign model
│   │   └── auth/
│   │       ├── jwt.go                  # JWT handling
│   │       └── middleware.go           # Auth middleware
│   ├── db/
│   │   └── migrations/
│   │       ├── 001_initial.sql         # Initial schema
│   │       ├── 002_indices.sql         # Indices
│   │       └── 003_audit.sql           # Audit tables
│   ├── Dockerfile
│   ├── Makefile
│   ├── go.mod
│   ├── go.sum
│   └── README.md
│
└── deploy/                              # Deployment manifests
    ├── kubernetes/
    │   ├── namespace.yaml
    │   ├── piko/
    │   │   ├── statefulset.yaml        # Piko StatefulSet
    │   │   ├── configmap.yaml          # Piko config
    │   │   ├── service.yaml            # Piko service
    │   │   └── hpa.yaml                # Horizontal autoscaling
    │   ├── control-plane/
    │   │   ├── deployment.yaml         # Control plane deployment
    │   │   ├── service.yaml            # Service
    │   │   ├── hpa.yaml                # Autoscaling
    │   │   └── secrets.yaml            # Secrets
    │   ├── mysql/
    │   │   ├── statefulset.yaml        # MySQL StatefulSet
    │   │   ├── service.yaml            # MySQL service
    │   │   ├── pvc.yaml                # Persistent volume
    │   │   └── configmap.yaml          # MySQL config
    │   ├── quickwit/
    │   │   ├── deployment.yaml         # Quickwit deployment
    │   │   ├── service.yaml            # Service
    │   │   └── pvc.yaml                # Storage
    │   └── ingress.yaml                # External ingress
    ├── docker/
    │   ├── agent.Dockerfile
    │   └── control-plane.Dockerfile
    └── scripts/
        ├── bootstrap.sh                 # Bootstrap script
        ├── deploy-agent.sh              # Agent deployment
        └── test-load.sh                 # Load testing
''')

# Dependencies
add_heading(doc4, '4. Dependencies and Version Pinning', 1)

add_heading(doc4, '4.1 Agent go.mod', 2)
add_code_block(doc4, '''module github.com/yourorg/vm-agent

go 1.21

require (
    github.com/spf13/cobra v1.8.0
    github.com/spf13/viper v1.18.2
    github.com/golang-jwt/jwt/v5 v5.2.0
    github.com/gorilla/websocket v1.5.1
    go.uber.org/zap v1.26.0
    golang.org/x/sys v0.16.0
)

// Piko client (to be embedded)
require (
    github.com/andydunstall/piko/client v0.8.0
)

// adnanh/webhook (to be embedded)
require (
    github.com/adnanh/webhook v2.8.1+incompatible
)

// linyows/probe (to be embedded)
require (
    github.com/linyows/probe v0.1.0
    gopkg.in/yaml.v3 v3.0.1
)
''')

add_heading(doc4, '4.2 Control Plane go.mod', 2)
add_code_block(doc4, '''module github.com/yourorg/control-plane

go 1.21

require (
    github.com/gin-gonic/gin v1.9.1
    github.com/golang-jwt/jwt/v5 v5.2.0
    github.com/spf13/viper v1.18.2
    github.com/go-sql-driver/mysql v1.7.1
    gorm.io/gorm v1.25.5
    gorm.io/driver/mysql v1.5.2
    github.com/prometheus/client_golang v1.18.0
    go.uber.org/zap v1.26.0
    golang.org/x/crypto v0.18.0
)

// Kubernetes client
require (
    k8s.io/client-go v0.29.0
    k8s.io/apimachinery v0.29.0
)

// Testing
require (
    github.com/stretchr/testify v1.8.4
    github.com/golang/mock v1.6.0
)
''')

# Build Instructions
add_heading(doc4, '5. Build Instructions', 1)

add_heading(doc4, '5.1 Agent Build', 2)
add_code_block(doc4, '''# Makefile for vm-agent
.PHONY: build test docker-build clean

BINARY_NAME=vm-agent
VERSION?=1.0.0
BUILD_DATE=$(shell date -u +"%Y-%m-%dT%H:%M:%SZ")
GIT_COMMIT=$(shell git rev-parse --short HEAD)
LDFLAGS=-ldflags "-X github.com/yourorg/vm-agent/internal/version.Version=$(VERSION) \\
                  -X github.com/yourorg/vm-agent/internal/version.BuildDate=$(BUILD_DATE) \\
                  -X github.com/yourorg/vm-agent/internal/version.GitCommit=$(GIT_COMMIT)"

build:
\tgo build $(LDFLAGS) -o bin/$(BINARY_NAME) cmd/agent/main.go

build-linux:
\tCGO_ENABLED=0 GOOS=linux GOARCH=amd64 go build $(LDFLAGS) -o bin/$(BINARY_NAME)-linux-amd64 cmd/agent/main.go

build-windows:
\tCGO_ENABLED=0 GOOS=windows GOARCH=amd64 go build $(LDFLAGS) -o bin/$(BINARY_NAME)-windows-amd64.exe cmd/agent/main.go

build-darwin:
\tCGO_ENABLED=0 GOOS=darwin GOARCH=amd64 go build $(LDFLAGS) -o bin/$(BINARY_NAME)-darwin-amd64 cmd/agent/main.go

build-all: build-linux build-windows build-darwin

test:
\tgo test -v -race -cover ./...

test-coverage:
\tgo test -v -race -coverprofile=coverage.out ./...
\tgo tool cover -html=coverage.out -o coverage.html

lint:
\tgolangci-lint run ./...

docker-build:
\tdocker build -t vm-agent:$(VERSION) -f Dockerfile .

clean:
\trm -rf bin/ coverage.out coverage.html

install:
\tgo install $(LDFLAGS) cmd/agent/main.go
''')

add_heading(doc4, '5.2 Control Plane Build', 2)
add_code_block(doc4, '''# Makefile for control-plane
.PHONY: build test docker-build migrate clean

BINARY_NAME=control-plane
VERSION?=1.0.0
BUILD_DATE=$(shell date -u +"%Y-%m-%dT%H:%M:%SZ")
GIT_COMMIT=$(shell git rev-parse --short HEAD)
LDFLAGS=-ldflags "-X github.com/yourorg/control-plane/internal/version.Version=$(VERSION) \\
                  -X github.com/yourorg/control-plane/internal/version.BuildDate=$(BUILD_DATE) \\
                  -X github.com/yourorg/control-plane/internal/version.GitCommit=$(GIT_COMMIT)"

build:
\tgo build $(LDFLAGS) -o bin/$(BINARY_NAME) cmd/server/main.go

test:
\tgo test -v -race -cover ./...

test-integration:
\tgo test -v -tags=integration ./...

lint:
\tgolangci-lint run ./...

docker-build:
\tdocker build -t control-plane:$(VERSION) -f Dockerfile .

migrate:
\t@echo "Running database migrations..."
\t@mysql -h $(DB_HOST) -u $(DB_USER) -p$(DB_PASSWORD) $(DB_NAME) < db/migrations/001_initial.sql
\t@mysql -h $(DB_HOST) -u $(DB_USER) -p$(DB_PASSWORD) $(DB_NAME) < db/migrations/002_indices.sql
\t@mysql -h $(DB_HOST) -u $(DB_USER) -p$(DB_PASSWORD) $(DB_NAME) < db/migrations/003_audit.sql

clean:
\trm -rf bin/ coverage.out coverage.html

run:
\tgo run cmd/server/main.go
''')

# Code Generation Tasks
add_heading(doc4, '6. Code Generation Tasks', 1)

add_heading(doc4, '6.1 Task 1: Agent Manager', 2)
add_paragraph(doc4, 'File: vm-agent/pkg/agent/manager.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• CLI parsing using cobra')
add_paragraph(doc4, '• Service management for all embedded components')
add_paragraph(doc4, '• Health monitoring loop')
add_paragraph(doc4, '• Graceful shutdown on SIGTERM')
add_paragraph(doc4, '• Component coordination')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Key Functions:', bold=True)
add_paragraph(doc4, '• NewManager(cfg *config.Config) (*Manager, error)')
add_paragraph(doc4, '• Run() error')
add_paragraph(doc4, '• RootCmd() *cobra.Command')
add_paragraph(doc4, '• InstallCmd() *cobra.Command')
add_paragraph(doc4, '• ConfigureCmd() *cobra.Command')
add_paragraph(doc4, '• RepairCmd() *cobra.Command')
add_paragraph(doc4, '• UpgradeCmd() *cobra.Command')
add_paragraph(doc4, '• UninstallCmd() *cobra.Command')
add_paragraph(doc4, '• StatusCmd() *cobra.Command')
add_paragraph(doc4, '• HealthCheck() health.Status')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Testing Requirements:', bold=True)
add_paragraph(doc4, '• Unit tests for all CLI commands')
add_paragraph(doc4, '• Mock all external dependencies')
add_paragraph(doc4, '• Test graceful shutdown')
add_paragraph(doc4, '• Test health check reporting')

add_heading(doc4, '6.2 Task 2: Lifecycle Commands', 2)
add_paragraph(doc4, 'Files:', bold=True)
add_paragraph(doc4, '• vm-agent/pkg/lifecycle/install.go')
add_paragraph(doc4, '• vm-agent/pkg/lifecycle/service_linux.go')
add_paragraph(doc4, '• vm-agent/pkg/lifecycle/service_windows.go')
add_paragraph(doc4, '• vm-agent/pkg/lifecycle/upgrade.go')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• systemd service integration for Linux')
add_paragraph(doc4, '• Windows Service integration using golang.org/x/sys/windows/svc')
add_paragraph(doc4, '• Self-upgrade with download, verify, backup, replace, rollback')
add_paragraph(doc4, '• Automatic rollback on upgrade failure')
add_paragraph(doc4, '• Version tracking in database')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Upgrade Flow:', bold=True)
add_paragraph(doc4, '1. Download new binary from control plane')
add_paragraph(doc4, '2. Verify SHA256 checksum')
add_paragraph(doc4, '3. Backup current binary')
add_paragraph(doc4, '4. Replace binary atomically')
add_paragraph(doc4, '5. Restart service')
add_paragraph(doc4, '6. Verify new version running')
add_paragraph(doc4, '7. Rollback if verification fails')

add_heading(doc4, '6.3 Task 3: Control Plane API', 2)
add_paragraph(doc4, 'Files: control-plane/pkg/api/*.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• REST API using Gin framework')
add_paragraph(doc4, '• JWT authentication middleware')
add_paragraph(doc4, '• Tenant isolation enforcement')
add_paragraph(doc4, '• OpenAPI/Swagger documentation')
add_paragraph(doc4, '• Rate limiting per tenant')
add_paragraph(doc4, '• Request/response logging')
add_paragraph(doc4, '• Prometheus metrics')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Endpoints to Implement:', bold=True)
add_code_block(doc4, '''POST   /api/v1/tenants
GET    /api/v1/tenants/:id
PUT    /api/v1/tenants/:id
DELETE /api/v1/tenants/:id
POST   /api/v1/tenants/:id/keys
POST   /api/v1/agents/register
GET    /api/v1/agents
GET    /api/v1/agents/:id
POST   /api/v1/workflows
GET    /api/v1/workflows/:id
POST   /api/v1/campaigns
GET    /api/v1/campaigns/:id
GET    /api/v1/audit
''')

add_heading(doc4, '6.4 Task 4: Agent Registration', 2)
add_paragraph(doc4, 'File: control-plane/pkg/agent/registration.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• One-time installation key validation')
add_paragraph(doc4, '• JWT token generation with 1-year expiry')
add_paragraph(doc4, '• Agent persistence to database')
add_paragraph(doc4, '• Automatic key expiration')
add_paragraph(doc4, '• Audit logging of all registrations')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Key Functions:', bold=True)
add_paragraph(doc4, '• GenerateInstallationKey(tenantID string, expiryHours int) (string, error)')
add_paragraph(doc4, '• ValidateInstallationKey(key string) (tenantID string, error)')
add_paragraph(doc4, '• RegisterAgent(key, agentID, version string) (token string, error)')
add_paragraph(doc4, '• MarkKeyUsed(keyHash string) error')

add_heading(doc4, '6.5 Task 5: MCP Server', 2)
add_paragraph(doc4, 'Files: control-plane/pkg/mcp/*.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• MCP protocol implementation')
add_paragraph(doc4, '• Tool handlers for workflow generation')
add_paragraph(doc4, '• LLM integration for natural language processing')
add_paragraph(doc4, '• Workflow validation')
add_paragraph(doc4, '• Integration with control plane API')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Tools to Implement:', bold=True)
add_paragraph(doc4, '• generate_workflow: Convert natural language to YAML workflow')
add_paragraph(doc4, '• validate_workflow: Validate workflow syntax and semantics')
add_paragraph(doc4, '• submit_workflow: Submit workflow to agents')
add_paragraph(doc4, '• list_agents: Query available agents')
add_paragraph(doc4, '• check_workflow_status: Get workflow execution status')

add_heading(doc4, '6.6 Task 6: Quickwit Integration', 2)
add_paragraph(doc4, 'Files: control-plane/pkg/audit/*.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• Quickwit HTTP client')
add_paragraph(doc4, '• Async log shipping with buffering')
add_paragraph(doc4, '• Index creation and management')
add_paragraph(doc4, '• Search query builder')
add_paragraph(doc4, '• Automatic retry on failure')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Log Schema:', bold=True)
add_code_block(doc4, '''{
  "timestamp": "2024-01-15T10:30:00Z",
  "tenant_id": "acme",
  "event_type": "workflow_execution",
  "actor": "admin@acme.com",
  "agent_id": "server-001",
  "action": "execute_workflow",
  "result": "success",
  "details": {
    "workflow_id": "wf-12345",
    "duration_ms": 1234
  },
  "ip_address": "203.0.113.1"
}''')

add_heading(doc4, '6.7 Task 7: Campaign Manager', 2)
add_paragraph(doc4, 'Files: control-plane/pkg/campaign/*.go', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• Phased rollout (canary, pilot, waves)')
add_paragraph(doc4, '• Success/failure tracking')
add_paragraph(doc4, '• Automatic rollback on threshold breach')
add_paragraph(doc4, '• Progress reporting')
add_paragraph(doc4, '• Concurrent workflow execution with limits')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Phase Configuration:', bold=True)
add_code_block(doc4, '''{
  "phases": [
    {
      "name": "canary",
      "percentage": 1,
      "success_threshold": 95,
      "wait_minutes": 30
    },
    {
      "name": "pilot",
      "percentage": 10,
      "success_threshold": 98,
      "wait_minutes": 60
    },
    {
      "name": "wave1",
      "percentage": 30,
      "success_threshold": 98,
      "wait_minutes": 30
    },
    {
      "name": "wave2",
      "percentage": 100,
      "success_threshold": 99,
      "wait_minutes": 0
    }
  ],
  "rollback_threshold": 90
}''')

add_heading(doc4, '6.8 Task 8: Database Migrations', 2)
add_paragraph(doc4, 'Files: control-plane/db/migrations/*.sql', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• All tables with tenant_id for multi-tenancy')
add_paragraph(doc4, '• Proper indices for query performance')
add_paragraph(doc4, '• Foreign key constraints')
add_paragraph(doc4, '• Audit trail tables')
add_paragraph(doc4, '• Migration versioning')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Tables to Create:', bold=True)
add_paragraph(doc4, '• tenants')
add_paragraph(doc4, '• agents')
add_paragraph(doc4, '• installation_keys')
add_paragraph(doc4, '• workflows')
add_paragraph(doc4, '• campaigns')
add_paragraph(doc4, '• audit_logs')

add_heading(doc4, '6.9 Task 9: Kubernetes Manifests', 2)
add_paragraph(doc4, 'Files: deploy/kubernetes/*.yaml', bold=True)
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• Piko StatefulSet with HPA')
add_paragraph(doc4, '• Control plane Deployment with HPA')
add_paragraph(doc4, '• MySQL StatefulSet with persistent storage')
add_paragraph(doc4, '• Quickwit Deployment')
add_paragraph(doc4, '• Services for all components')
add_paragraph(doc4, '• Ingress with TLS termination')
add_paragraph(doc4, '• ConfigMaps and Secrets')
add_paragraph(doc4, '')
add_paragraph(doc4, 'HPA Configuration:', bold=True)
add_paragraph(doc4, '• Piko: Scale based on active connections (target: 10K/pod)')
add_paragraph(doc4, '• Control Plane: Scale based on CPU (target: 70%) and memory (target: 80%)')
add_paragraph(doc4, '• Min replicas: 3')
add_paragraph(doc4, '• Max replicas: 50')

add_heading(doc4, '6.10 Task 10: Docker Images', 2)
add_paragraph(doc4, 'Files: Dockerfile.*, bold=True')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Requirements:', bold=True)
add_paragraph(doc4, '• Multi-stage builds for minimal image size')
add_paragraph(doc4, '• Target size: <50MB per image')
add_paragraph(doc4, '• Non-root user execution')
add_paragraph(doc4, '• Health check endpoints')
add_paragraph(doc4, '• Build argument for version')
add_paragraph(doc4, '')
add_paragraph(doc4, 'Example Dockerfile:', bold=True)
add_code_block(doc4, '''# Agent Dockerfile
FROM golang:1.21-alpine AS builder

WORKDIR /build
COPY go.mod go.sum ./
RUN go mod download

COPY . .
RUN CGO_ENABLED=0 go build -ldflags="-s -w" -o vm-agent cmd/agent/main.go

FROM alpine:3.19
RUN apk --no-cache add ca-certificates

WORKDIR /app
COPY --from=builder /build/vm-agent /app/vm-agent

RUN addgroup -g 1000 vm-agent && \\
    adduser -u 1000 -G vm-agent -s /bin/sh -D vm-agent

USER vm-agent

ENTRYPOINT ["/app/vm-agent"]
CMD ["run"]
''')

# Configuration Examples
add_heading(doc4, '7. Configuration Examples', 1)

add_heading(doc4, '7.1 Agent Configuration', 2)
add_code_block(doc4, '''# /etc/vm-agent/config.yaml
agent:
  id: "server-001"
  tenant_id: "acme"
  control_plane_url: "https://control-plane.example.com"

piko:
  server_url: "https://piko.example.com"
  endpoint: "tenant-acme/server-001"
  reconnect:
    initial_delay: 1s
    max_delay: 60s
    multiplier: 2.0

webhook:
  listen_addr: "0.0.0.0"
  port: 9999
  tls_enabled: false

probe:
  work_dir: "/var/lib/vm-agent/work"
  default_timeout: 300s
  max_concurrent: 5

health:
  check_interval: 30s
  report_interval: 300s
  report_url: "https://control-plane.example.com/api/v1/agents/health"
''')

add_heading(doc4, '7.2 Control Plane Configuration', 2)
add_code_block(doc4, '''# config.yaml
server:
  listen_addr: "0.0.0.0"
  port: 8080
  tls_enabled: true
  cert_file: "/etc/tls/cert.pem"
  key_file: "/etc/tls/key.pem"

database:
  host: "mysql"
  port: 3306
  username: "control_plane"
  password: "${DB_PASSWORD}"
  database: "vm_manager"
  max_connections: 100
  max_idle_connections: 10
  connection_lifetime: 3600s

piko:
  server_url: "http://piko:8000"
  admin_token: "${PIKO_ADMIN_TOKEN}"

jwt:
  secret: "${JWT_SECRET}"
  expiry_hours: 8760  # 1 year

quickwit:
  url: "http://quickwit:7280"
  index: "vm-agent-audit-logs"

mcp:
  enabled: true
  llm_api_key: "${LLM_API_KEY}"
  llm_model: "claude-3-sonnet"
''')

# Testing Requirements
add_heading(doc4, '8. Testing Requirements', 1)

add_paragraph(doc4, 'Unit Tests:', bold=True)
add_paragraph(doc4, '• Coverage target: >80%')
add_paragraph(doc4, '• Mock all external dependencies')
add_paragraph(doc4, '• Test all error paths')
add_paragraph(doc4, '• Use table-driven tests')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Integration Tests:', bold=True)
add_paragraph(doc4, '• API endpoint tests')
add_paragraph(doc4, '• Database integration tests')
add_paragraph(doc4, '• Piko client integration')
add_paragraph(doc4, '• Workflow execution tests')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Load Tests:', bold=True)
add_paragraph(doc4, '• 100K concurrent agents')
add_paragraph(doc4, '• 1M agent registration simulation')
add_paragraph(doc4, '• Campaign rollout at scale')
add_paragraph(doc4, '• API throughput: 10K req/sec')

# Coding Standards
add_heading(doc4, '9. Coding Standards', 1)

add_paragraph(doc4, 'Go Style Guide:', bold=True)
add_paragraph(doc4, '• Follow https://golang.org/doc/effective_go')
add_paragraph(doc4, '• Use gofmt for formatting')
add_paragraph(doc4, '• Use golangci-lint for linting')
add_paragraph(doc4, '• Package names: lowercase, no underscores')
add_paragraph(doc4, '• Exported functions: Start with capital letter')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Error Handling:', bold=True)
add_paragraph(doc4, '• Always check errors')
add_paragraph(doc4, '• Wrap errors with context using fmt.Errorf')
add_paragraph(doc4, '• Log errors at appropriate level')
add_paragraph(doc4, '• Return errors, don\'t panic')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Logging Standards:', bold=True)
add_paragraph(doc4, '• Use structured logging (zap)')
add_paragraph(doc4, '• Include context (tenant_id, agent_id, etc.)')
add_paragraph(doc4, '• Log levels: DEBUG, INFO, WARN, ERROR')
add_paragraph(doc4, '• No sensitive data in logs')

# Deliverables Checklist
add_heading(doc4, '10. Deliverables Checklist', 1)

add_paragraph(doc4, '□ Agent binary (Linux, Windows, macOS)', bold=True)
add_paragraph(doc4, '□ Control plane binary')
add_paragraph(doc4, '□ Docker images (<50MB)')
add_paragraph(doc4, '□ Kubernetes manifests with HPA')
add_paragraph(doc4, '□ Database migrations')
add_paragraph(doc4, '□ OpenAPI documentation')
add_paragraph(doc4, '□ Unit tests (>80% coverage)')
add_paragraph(doc4, '□ Integration tests')
add_paragraph(doc4, '□ Load test scripts')
add_paragraph(doc4, '□ Deployment scripts')
add_paragraph(doc4, '□ Operations runbook')
add_paragraph(doc4, '□ README files')

# Timeline
add_heading(doc4, '11. Implementation Timeline', 1)

add_paragraph(doc4, 'Phase 1: Core Agent (4 weeks)', bold=True)
add_paragraph(doc4, '• Week 1: Project setup, agent manager, configuration')
add_paragraph(doc4, '• Week 2: Piko client, webhook server, probe executor')
add_paragraph(doc4, '• Week 3: Lifecycle management (install, upgrade, repair)')
add_paragraph(doc4, '• Week 4: Health monitoring, testing')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Phase 2: Control Plane (4 weeks)', bold=True)
add_paragraph(doc4, '• Week 1: API server, database schema, migrations')
add_paragraph(doc4, '• Week 2: Tenant management, agent registration')
add_paragraph(doc4, '• Week 3: Workflow management, campaign manager')
add_paragraph(doc4, '• Week 4: Testing, documentation')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Phase 3: MCP Integration (2 weeks)', bold=True)
add_paragraph(doc4, '• Week 1: MCP server, tool definitions')
add_paragraph(doc4, '• Week 2: Workflow generation, validation, testing')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Phase 4: Testing & Documentation (2 weeks)', bold=True)
add_paragraph(doc4, '• Week 1: Integration tests, load tests')
add_paragraph(doc4, '• Week 2: Documentation, deployment guides')

# Success Criteria
add_heading(doc4, '12. Success Criteria', 1)

add_paragraph(doc4, 'Performance:', bold=True)
add_paragraph(doc4, '• Agent binary: <20MB')
add_paragraph(doc4, '• Agent memory: <4MB per instance')
add_paragraph(doc4, '• Agent registration: <500ms p95')
add_paragraph(doc4, '• Workflow latency: <200ms p95')
add_paragraph(doc4, '• API throughput: 10,000 req/sec')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Scalability:', bold=True)
add_paragraph(doc4, '• Support 1M+ agents')
add_paragraph(doc4, '• Agent upgrade across 1M VMs: <72 hours')
add_paragraph(doc4, '• Upgrade success rate: >99.8%')
add_paragraph(doc4, '')

add_paragraph(doc4, 'Reliability:', bold=True)
add_paragraph(doc4, '• System availability: 99.95%')
add_paragraph(doc4, '• Zero-downtime upgrades')
add_paragraph(doc4, '• Automatic rollback on failure')
add_paragraph(doc4, '• Complete audit trail')

# Save Document 4
doc4.save('doc4_claude_code_spec.docx')
print("Document 4 completed: doc4_claude_code_spec.docx")

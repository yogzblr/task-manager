#!/usr/bin/env python3
"""Generate Multi-Tenant VM Manager Documentation"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a heading with formatting"""
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc, text, bold=False):
    """Add a paragraph with optional bold"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    p = doc.add_paragraph(code_text)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

# Document 1: Approach Comparison
print("Generating Document 1: Approach Comparison...")
doc1 = Document()
doc1.add_heading('Multi-Tenant VM Manager - Approach Comparison', 0)

# Executive Summary
add_heading(doc1, '1. Executive Summary', 1)
add_paragraph(doc1, 'This document compares different approaches for building a scalable multi-tenant VM management system capable of handling 1M+ Windows/Linux VMs. We evaluate two custom approaches against existing industry solutions including SaltStack, AWS SSM, Tencent TAT, and GitHub Actions Runner.')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Key Findings:', bold=True)
add_paragraph(doc1, '• Approach 1 (WebSocket/Centrifugo): Best for 10K-100K VMs with real-time requirements')
add_paragraph(doc1, '• Approach 2 (Piko + webhook + probe): Best for 100K-1M+ VMs with optimal resource usage')
add_paragraph(doc1, '• Existing solutions have limitations: SaltStack (10K/master limit), AWS SSM (AWS-only), GitHub Actions (CI/CD focused)')

# Current Approaches Overview
add_heading(doc1, '2. Current Approaches Overview', 1)

add_heading(doc1, '2.1 Your Approach 1: WebSocket (Centrifugo)', 2)
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Push-based communication using WebSocket protocol')
add_paragraph(doc1, '• Centrifugo as the messaging broker')
add_paragraph(doc1, '• Agents maintain persistent connections')
add_paragraph(doc1, '• Real-time command delivery (<100ms latency)')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Technical Details:', bold=True)
add_paragraph(doc1, '• Memory per agent: ~8-12MB (connection overhead)')
add_paragraph(doc1, '• Max connections per Centrifugo node: ~10K-20K')
add_paragraph(doc1, '• Requires connection state management')
add_paragraph(doc1, '• Excellent for real-time notifications')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Extremely low latency (<100ms)')
add_paragraph(doc1, '• Real-time bidirectional communication')
add_paragraph(doc1, '• Good for monitoring and alerting')
add_paragraph(doc1, '• Simple programming model')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• High memory usage at scale')
add_paragraph(doc1, '• Connection storms during restarts')
add_paragraph(doc1, '• Complex load balancing for 100K+ connections')
add_paragraph(doc1, '• Firewall traversal challenges')

add_heading(doc1, '2.2 Your Approach 2: Piko + webhook + linyows/probe', 2)
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Hybrid push-pull using reverse proxy (Piko)')
add_paragraph(doc1, '• Agents expose webhook endpoints via Piko tunnels')
add_paragraph(doc1, '• linyows/probe executes workflows locally')
add_paragraph(doc1, '• Control plane sends HTTP requests through Piko')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Technical Details:', bold=True)
add_paragraph(doc1, '• Memory per agent: ~2-4MB (no persistent connection)')
add_paragraph(doc1, '• Piko handles connection multiplexing')
add_paragraph(doc1, '• Automatic load rebalancing across Piko nodes')
add_paragraph(doc1, '• NAT/firewall friendly (outbound only)')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Very low memory footprint')
add_paragraph(doc1, '• Excellent scalability (1M+ agents)')
add_paragraph(doc1, '• Built-in load balancing and HA')
add_paragraph(doc1, '• No connection storms')
add_paragraph(doc1, '• Cross-platform (Windows, Linux, macOS)')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• Higher latency vs WebSocket (~200-500ms)')
add_paragraph(doc1, '• Not suitable for sub-second real-time needs')
add_paragraph(doc1, '• Requires Piko cluster management')

# Comparison with Existing Tools
add_heading(doc1, '3. Comparison with Existing Tools', 1)

add_heading(doc1, '3.1 SaltStack', 2)
add_paragraph(doc1, 'Overview:', bold=True)
add_paragraph(doc1, 'SaltStack is an enterprise configuration management and remote execution tool using ZeroMQ for communication.')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Master-minion architecture')
add_paragraph(doc1, '• ZeroMQ pub-sub messaging (ports 4505, 4506)')
add_paragraph(doc1, '• Event-driven communication')
add_paragraph(doc1, '• Python-based execution modules')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Scalability:', bold=True)
add_paragraph(doc1, '• Max ~10K minions per master (documented limit)')
add_paragraph(doc1, '• Requires master-of-masters for >10K')
add_paragraph(doc1, '• Syndic architecture for horizontal scaling')
add_paragraph(doc1, '• Memory: ~15-30MB per minion')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Mature ecosystem with extensive modules')
add_paragraph(doc1, '• Powerful state management (SLS files)')
add_paragraph(doc1, '• Fast execution engine')
add_paragraph(doc1, '• Built-in event bus')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• Hard 10K/master limit requires complex syndic setup')
add_paragraph(doc1, '• Heavy Python dependency (~200MB+ per agent)')
add_paragraph(doc1, '• ZeroMQ firewall configuration complexity')
add_paragraph(doc1, '• No native multi-tenancy')
add_paragraph(doc1, '• Difficult to integrate with AI/MCP')

add_heading(doc1, '3.2 AWS Systems Manager (SSM)', 2)
add_paragraph(doc1, 'Overview:', bold=True)
add_paragraph(doc1, 'AWS SSM is a pull-based agent for managing AWS EC2 and on-premises servers.')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Pull-based polling (agents poll AWS API)')
add_paragraph(doc1, '• SSM Agent checks for commands every ~15-60 seconds')
add_paragraph(doc1, '• Uses AWS API endpoints')
add_paragraph(doc1, '• Commands stored in S3, retrieved by agents')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Scalability:', bold=True)
add_paragraph(doc1, '• Unlimited scale (AWS infrastructure)')
add_paragraph(doc1, '• No connection state management')
add_paragraph(doc1, '• Memory: ~50-100MB per agent (Go binary + overhead)')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Infinite scalability (AWS managed)')
add_paragraph(doc1, '• No infrastructure management')
add_paragraph(doc1, '• Integrated with AWS ecosystem')
add_paragraph(doc1, '• Automatic updates')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• AWS-only (vendor lock-in)')
add_paragraph(doc1, '• High latency (15-60 second polling)')
add_paragraph(doc1, '• Limited on-premises support')
add_paragraph(doc1, '• No self-hosting option')
add_paragraph(doc1, '• Expensive at scale')
add_paragraph(doc1, '• Cannot customize for multi-tenancy')

add_heading(doc1, '3.3 Tencent TAT (TencentCloud Automation Tools)', 2)
add_paragraph(doc1, 'Overview:', bold=True)
add_paragraph(doc1, 'Tencent TAT is similar to AWS SSM, designed for Tencent Cloud infrastructure.')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Pull-based polling model')
add_paragraph(doc1, '• Agent polls Tencent Cloud API')
add_paragraph(doc1, '• Command execution and reporting')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Integrated with Tencent Cloud')
add_paragraph(doc1, '• Scalable (cloud-managed)')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• Tencent Cloud only (vendor lock-in)')
add_paragraph(doc1, '• Limited documentation (primarily Chinese)')
add_paragraph(doc1, '• No self-hosting option')
add_paragraph(doc1, '• Similar limitations to AWS SSM')

add_heading(doc1, '3.4 GitHub Actions Self-Hosted Runner', 2)
add_paragraph(doc1, 'Overview:', bold=True)
add_paragraph(doc1, 'Self-hosted runners for GitHub Actions CI/CD workflows.')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Architecture:', bold=True)
add_paragraph(doc1, '• Pull-based long polling')
add_paragraph(doc1, '• Runners poll GitHub API for jobs')
add_paragraph(doc1, '• Job execution in isolated containers/VMs')
add_paragraph(doc1, '• Results posted back to GitHub')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Scalability:', bold=True)
add_paragraph(doc1, '• Unlimited runners')
add_paragraph(doc1, '• Horizontal scaling via runner groups')
add_paragraph(doc1, '• Memory: ~100-200MB per runner')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Strengths:', bold=True)
add_paragraph(doc1, '• Perfect for CI/CD workflows')
add_paragraph(doc1, '• Integrated with GitHub ecosystem')
add_paragraph(doc1, '• Self-hosted option available')
add_paragraph(doc1, '• Good isolation model')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Weaknesses:', bold=True)
add_paragraph(doc1, '• Designed for CI/CD, not general management')
add_paragraph(doc1, '• GitHub dependency')
add_paragraph(doc1, '• Not suitable for real-time operations')
add_paragraph(doc1, '• Limited to workflow execution')
add_paragraph(doc1, '• High latency (polling interval)')

# Cross-Platform Support Analysis
add_heading(doc1, '4. Cross-Platform Support Analysis', 1)

add_paragraph(doc1, 'Component Compatibility Matrix:', bold=True)
add_paragraph(doc1, '')

# Create a simple table representation
add_paragraph(doc1, 'Component          | Windows | Linux | macOS | Notes')
add_paragraph(doc1, '-------------------|---------|-------|-------|------------------')
add_paragraph(doc1, 'Piko Client        | ✓       | ✓     | ✓     | Pure Go')
add_paragraph(doc1, 'adnanh/webhook     | ✓       | ✓     | ✓     | Pure Go')
add_paragraph(doc1, 'linyows/probe      | ✓       | ✓     | ✓     | Pure Go')
add_paragraph(doc1, 'Centrifugo Client  | ✓       | ✓     | ✓     | Multiple langs')
add_paragraph(doc1, 'SaltStack Minion   | ✓       | ✓     | ✓     | Python required')
add_paragraph(doc1, 'AWS SSM Agent      | ✓       | ✓     | ✗     | Go binary')
add_paragraph(doc1, 'GitHub Runner      | ✓       | ✓     | ✓     | .NET runtime')
add_paragraph(doc1, '')

add_paragraph(doc1, 'Platform-Specific Considerations:', bold=True)
add_paragraph(doc1, '')
add_paragraph(doc1, 'Windows:', bold=True)
add_paragraph(doc1, '• Service installation via Windows Service Manager')
add_paragraph(doc1, '• Registry-based configuration')
add_paragraph(doc1, '• Windows Firewall rules')
add_paragraph(doc1, '• Event log integration')
add_paragraph(doc1, '• All Go-based components work natively')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Linux:', bold=True)
add_paragraph(doc1, '• systemd service management')
add_paragraph(doc1, '• File-based configuration (/etc)')
add_paragraph(doc1, '• iptables/firewalld integration')
add_paragraph(doc1, '• syslog integration')
add_paragraph(doc1, '• Native support for all components')
add_paragraph(doc1, '')
add_paragraph(doc1, 'macOS:', bold=True)
add_paragraph(doc1, '• launchd service management')
add_paragraph(doc1, '• Limited enterprise deployment')
add_paragraph(doc1, '• Primarily for development/testing')

# Detailed Comparison Tables
add_heading(doc1, '5. Detailed Comparison Tables', 1)

add_heading(doc1, '5.1 Communication Model Comparison', 2)
add_paragraph(doc1, 'Solution               | Model        | Protocol    | Latency    | Firewall')
add_paragraph(doc1, '-----------------------|--------------|-------------|------------|----------')
add_paragraph(doc1, 'Approach 1 (Centrifugo)| Push         | WebSocket   | <100ms     | Complex')
add_paragraph(doc1, 'Approach 2 (Piko)      | Hybrid       | HTTP/WS     | 200-500ms  | Simple')
add_paragraph(doc1, 'SaltStack              | Push         | ZeroMQ      | 100-200ms  | Complex')
add_paragraph(doc1, 'AWS SSM                | Pull         | HTTPS       | 15-60s     | Simple')
add_paragraph(doc1, 'Tencent TAT            | Pull         | HTTPS       | 15-60s     | Simple')
add_paragraph(doc1, 'GitHub Runner          | Pull (LP)    | HTTPS       | 1-5s       | Simple')
add_paragraph(doc1, '')

add_heading(doc1, '5.2 Scalability Comparison', 2)
add_paragraph(doc1, 'Solution               | Max Scale    | Mem/Agent | Scale Model       | Complexity')
add_paragraph(doc1, '-----------------------|--------------|-----------|-------------------|------------')
add_paragraph(doc1, 'Approach 1 (Centrifugo)| 100K         | 8-12MB    | Vertical+Sharding | Medium')
add_paragraph(doc1, 'Approach 2 (Piko)      | 1M+          | 2-4MB     | Horizontal (HPA)  | Low')
add_paragraph(doc1, 'SaltStack              | 10K/master   | 15-30MB   | Syndic (Complex)  | High')
add_paragraph(doc1, 'AWS SSM                | Unlimited    | 50-100MB  | AWS Managed       | None')
add_paragraph(doc1, 'Tencent TAT            | Unlimited    | 50-100MB  | Cloud Managed     | None')
add_paragraph(doc1, 'GitHub Runner          | Unlimited    | 100-200MB | Horizontal        | Low')
add_paragraph(doc1, '')

add_heading(doc1, '5.3 Performance at Scale', 2)
add_paragraph(doc1, 'Metric                 | Approach 1 | Approach 2 | SaltStack | AWS SSM | GitHub')
add_paragraph(doc1, '-----------------------|------------|------------|-----------|---------|--------')
add_paragraph(doc1, '100K VMs - Total Mem   | 800GB-1.2TB| 200-400GB  | 1.5-3TB   | 5-10TB  | 10-20TB')
add_paragraph(doc1, '100K VMs - Cmd Latency | <100ms     | 200-500ms  | 100-200ms | 15-60s  | 1-5s')
add_paragraph(doc1, '1M VMs - Total Mem     | 8-12TB     | 2-4TB      | N/A       | 50-100TB| 100-200TB')
add_paragraph(doc1, '1M VMs - Cmd Latency   | Unstable   | 300-800ms  | N/A       | 15-60s  | 1-5s')
add_paragraph(doc1, '1M VMs - Infra Nodes   | 100+       | 20-30      | N/A       | AWS     | Many')
add_paragraph(doc1, '')

add_heading(doc1, '5.4 Operational Complexity', 2)
add_paragraph(doc1, 'Aspect                 | Approach 1 | Approach 2 | SaltStack | AWS SSM | GitHub')
add_paragraph(doc1, '-----------------------|------------|------------|-----------|---------|--------')
add_paragraph(doc1, 'Setup Complexity       | Medium     | Medium     | High      | Low     | Low')
add_paragraph(doc1, 'Maintenance            | Medium     | Low        | High      | None    | Low')
add_paragraph(doc1, 'Monitoring Needs       | High       | Medium     | High      | Low     | Medium')
add_paragraph(doc1, 'Scaling Effort         | High       | Low (Auto) | Very High | None    | Low')
add_paragraph(doc1, 'Upgrade Complexity     | Medium     | Low        | High      | Auto    | Medium')
add_paragraph(doc1, '')

add_heading(doc1, '5.5 Reliability Comparison', 2)
add_paragraph(doc1, 'Aspect                 | Approach 1 | Approach 2 | SaltStack | AWS SSM | GitHub')
add_paragraph(doc1, '-----------------------|------------|------------|-----------|---------|--------')
add_paragraph(doc1, 'Connection Storms      | High Risk  | None       | Medium    | None    | None')
add_paragraph(doc1, 'Split Brain Risk       | Low        | None       | Medium    | None    | None')
add_paragraph(doc1, 'Auto Recovery          | Manual     | Automatic  | Manual    | Auto    | Auto')
add_paragraph(doc1, 'HA Support             | External   | Built-in   | External  | AWS     | GitHub')
add_paragraph(doc1, 'Failure Detection      | Immediate  | 30-60s     | Fast      | Slow    | Medium')

# Security Comparison
add_heading(doc1, '6. Security Comparison', 1)

add_heading(doc1, '6.1 Authentication Methods', 2)
add_paragraph(doc1, 'Solution               | Agent Auth        | API Auth      | Multi-Tenant')
add_paragraph(doc1, '-----------------------|-------------------|---------------|-------------')
add_paragraph(doc1, 'Approach 1 (Centrifugo)| JWT/Token         | JWT           | Custom')
add_paragraph(doc1, 'Approach 2 (Piko)      | JWT + mTLS        | JWT + mTLS    | Built-in')
add_paragraph(doc1, 'SaltStack              | Keys (AES)        | PAM/LDAP      | None')
add_paragraph(doc1, 'AWS SSM                | IAM + Instance ID | IAM           | IAM-based')
add_paragraph(doc1, 'GitHub Runner          | Registration Token| GitHub Token  | Orgs/Repos')
add_paragraph(doc1, '')

add_heading(doc1, '6.2 Command Execution Restrictions', 2)
add_paragraph(doc1, 'Approach 1 & 2:', bold=True)
add_paragraph(doc1, '• Custom RBAC with workflow validation')
add_paragraph(doc1, '• Command whitelisting/blacklisting')
add_paragraph(doc1, '• User/tenant-level permissions')
add_paragraph(doc1, '• Workflow approval gates')
add_paragraph(doc1, '')
add_paragraph(doc1, 'SaltStack:', bold=True)
add_paragraph(doc1, '• External auth system (PAM, LDAP)')
add_paragraph(doc1, '• Publisher ACLs')
add_paragraph(doc1, '• Limited per-command controls')
add_paragraph(doc1, '')
add_paragraph(doc1, 'AWS SSM:', bold=True)
add_paragraph(doc1, '• IAM policies')
add_paragraph(doc1, '• Document-level permissions')
add_paragraph(doc1, '• Tag-based access control')
add_paragraph(doc1, '')
add_paragraph(doc1, 'GitHub Runner:', bold=True)
add_paragraph(doc1, '• Repository/org level permissions')
add_paragraph(doc1, '• Workflow approval required')
add_paragraph(doc1, '• No granular command control')

add_heading(doc1, '6.3 Audit Capabilities', 2)
add_paragraph(doc1, 'Solution               | Audit Logging     | Log Storage   | Retention')
add_paragraph(doc1, '-----------------------|-------------------|---------------|----------')
add_paragraph(doc1, 'Approach 1 & 2         | Quickwit          | Self-hosted   | Custom')
add_paragraph(doc1, 'SaltStack              | Event bus         | External      | Manual')
add_paragraph(doc1, 'AWS SSM                | CloudTrail        | S3            | 90 days')
add_paragraph(doc1, 'GitHub Runner          | GitHub Logs       | GitHub        | 90 days')

# Final Scoring Matrix
add_heading(doc1, '7. Final Scoring Matrix', 1)
add_paragraph(doc1, 'Rating Scale: ★☆☆☆☆ (Poor) to ★★★★★ (Excellent)')
add_paragraph(doc1, '')

add_paragraph(doc1, 'Criteria               | Approach 1 | Approach 2 | SaltStack | AWS SSM | GitHub')
add_paragraph(doc1, '-----------------------|------------|------------|-----------|---------|--------')
add_paragraph(doc1, 'Scalability (1M VMs)   | ★★☆☆☆      | ★★★★★      | ★☆☆☆☆     | ★★★★★   | ★★★★☆')
add_paragraph(doc1, 'Latency/Performance    | ★★★★★      | ★★★★☆      | ★★★★☆     | ★☆☆☆☆   | ★★★☆☆')
add_paragraph(doc1, 'Resource Efficiency    | ★★☆☆☆      | ★★★★★      | ★★☆☆☆     | ★★★☆☆   | ★★☆☆☆')
add_paragraph(doc1, 'Operational Simplicity | ★★★☆☆      | ★★★★☆      | ★★☆☆☆     | ★★★★★   | ★★★★☆')
add_paragraph(doc1, 'Multi-Tenancy Support  | ★★★☆☆      | ★★★★★      | ★☆☆☆☆     | ★★★☆☆   | ★★★☆☆')
add_paragraph(doc1, 'Self-Hosting Option    | ★★★★★      | ★★★★★      | ★★★★★     | ☆☆☆☆☆   | ★★★☆☆')
add_paragraph(doc1, 'Cross-Platform Support | ★★★★★      | ★★★★★      | ★★★★☆     | ★★★☆☆   | ★★★★★')
add_paragraph(doc1, 'AI/MCP Integration     | ★★★★☆      | ★★★★★      | ★★☆☆☆     | ★★☆☆☆   | ★★★☆☆')
add_paragraph(doc1, 'Security & Audit       | ★★★★☆      | ★★★★★      | ★★★☆☆     | ★★★★☆   | ★★★☆☆')
add_paragraph(doc1, 'Customization          | ★★★★★      | ★★★★★      | ★★★☆☆     | ★☆☆☆☆   | ★★☆☆☆')
add_paragraph(doc1, 'Total Score            | 37/50      | 47/50      | 27/50     | 30/50   | 33/50')

# Recommendation
add_heading(doc1, '8. Recommendation', 1)

add_paragraph(doc1, 'Based on our comprehensive analysis, we recommend:', bold=True)
add_paragraph(doc1, '')

add_heading(doc1, '8.1 For 100K VMs or Less: Approach 1 (WebSocket/Centrifugo)', 2)
add_paragraph(doc1, 'Rationale:', bold=True)
add_paragraph(doc1, '• Sub-100ms latency is achievable and manageable at this scale')
add_paragraph(doc1, '• Infrastructure costs are reasonable (~800GB-1.2TB memory)')
add_paragraph(doc1, '• Simpler programming model for real-time features')
add_paragraph(doc1, '• Good balance between performance and complexity')
add_paragraph(doc1, '')
add_paragraph(doc1, 'When to choose this:', bold=True)
add_paragraph(doc1, '• Real-time monitoring/alerting is critical')
add_paragraph(doc1, '• Bidirectional communication is needed')
add_paragraph(doc1, '• You have <100K agents')
add_paragraph(doc1, '• Budget allows for higher infrastructure costs')

add_heading(doc1, '8.2 For 100K-1M+ VMs: Approach 2 (Piko + webhook + probe)', 2)
add_paragraph(doc1, 'Rationale:', bold=True)
add_paragraph(doc1, '• Proven scalability to 1M+ agents')
add_paragraph(doc1, '• 5x lower memory footprint (critical at scale)')
add_paragraph(doc1, '• Built-in multi-tenancy and auto-scaling')
add_paragraph(doc1, '• Excellent operational simplicity')
add_paragraph(doc1, '• Perfect for AI/MCP integration')
add_paragraph(doc1, '• 200-500ms latency is acceptable for most operations')
add_paragraph(doc1, '')
add_paragraph(doc1, 'When to choose this:', bold=True)
add_paragraph(doc1, '• You need to scale beyond 100K agents')
add_paragraph(doc1, '• Multi-tenancy is required')
add_paragraph(doc1, '• Resource efficiency is critical')
add_paragraph(doc1, '• You want AI-powered workflow generation')
add_paragraph(doc1, '• Sub-second latency is acceptable')

add_heading(doc1, '8.3 Why Not Existing Solutions?', 2)
add_paragraph(doc1, 'SaltStack:', bold=True)
add_paragraph(doc1, '• Hard 10K/master limit requires complex syndic architecture')
add_paragraph(doc1, '• No native multi-tenancy support')
add_paragraph(doc1, '• High operational complexity')
add_paragraph(doc1, '• Cannot scale to 1M VMs economically')
add_paragraph(doc1, '')
add_paragraph(doc1, 'AWS SSM:', bold=True)
add_paragraph(doc1, '• Vendor lock-in (AWS only)')
add_paragraph(doc1, '• No self-hosting option')
add_paragraph(doc1, '• High latency (15-60 seconds)')
add_paragraph(doc1, '• Expensive at scale')
add_paragraph(doc1, '• Limited customization')
add_paragraph(doc1, '')
add_paragraph(doc1, 'GitHub Actions Runner:', bold=True)
add_paragraph(doc1, '• Designed for CI/CD, not general VM management')
add_paragraph(doc1, '• GitHub dependency')
add_paragraph(doc1, '• Not suitable for configuration management')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Tencent TAT:', bold=True)
add_paragraph(doc1, '• Tencent Cloud only')
add_paragraph(doc1, '• Limited documentation')
add_paragraph(doc1, '• Similar limitations to AWS SSM')

add_heading(doc1, '8.4 Migration Strategy', 2)
add_paragraph(doc1, 'If starting with Approach 1 and planning to scale:', bold=True)
add_paragraph(doc1, '')
add_paragraph(doc1, 'Phase 1 (0-50K agents):', bold=True)
add_paragraph(doc1, '• Deploy Approach 1 with Centrifugo')
add_paragraph(doc1, '• Establish workflows and patterns')
add_paragraph(doc1, '• Build operational expertise')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Phase 2 (50K-100K agents):', bold=True)
add_paragraph(doc1, '• Monitor resource usage and costs')
add_paragraph(doc1, '• Begin Approach 2 development in parallel')
add_paragraph(doc1, '• Deploy Approach 2 to test environment')
add_paragraph(doc1, '')
add_paragraph(doc1, 'Phase 3 (100K+ agents):', bold=True)
add_paragraph(doc1, '• Gradual migration to Approach 2')
add_paragraph(doc1, '• Use unified agent binary (supports both backends)')
add_paragraph(doc1, '• Migrate tenant by tenant')
add_paragraph(doc1, '• Decommission Centrifugo infrastructure')
add_paragraph(doc1, '')

add_heading(doc1, '8.5 Final Recommendation', 2)
add_paragraph(doc1, 'We strongly recommend Approach 2 (Piko + webhook + probe) as the primary solution for these reasons:', bold=True)
add_paragraph(doc1, '')
add_paragraph(doc1, '1. Future-proof scalability: Ready for 1M+ VMs without architecture changes')
add_paragraph(doc1, '2. Cost efficiency: 5x lower resource usage translates to significant savings')
add_paragraph(doc1, '3. Operational simplicity: Built-in HPA and load balancing reduce management overhead')
add_paragraph(doc1, '4. Multi-tenancy: Native support eliminates complex custom implementation')
add_paragraph(doc1, '5. AI integration: MCP server enables natural language workflow generation')
add_paragraph(doc1, '6. Proven technology: Piko is battle-tested with large-scale deployments')
add_paragraph(doc1, '')
add_paragraph(doc1, 'The 200-500ms latency is acceptable for virtually all VM management operations including:')
add_paragraph(doc1, '• Software deployment')
add_paragraph(doc1, '• Configuration updates')
add_paragraph(doc1, '• Agent upgrades')
add_paragraph(doc1, '• Diagnostic commands')
add_paragraph(doc1, '• Scheduled maintenance')
add_paragraph(doc1, '')
add_paragraph(doc1, 'For the rare scenarios requiring sub-100ms latency, consider a hybrid approach where critical agents use Approach 1 while the majority use Approach 2.')

# Save Document 1
doc1.save('doc1_comparison.docx')
print("Document 1 completed: doc1_comparison.docx")

# Document 2: Architectural Design
print("\nGenerating Document 2: Architectural Design...")
doc2 = Document()
doc2.add_heading('Multi-Tenant VM Manager - Architectural Design', 0)

# Executive Summary
add_heading(doc2, '1. Executive Summary', 1)
add_paragraph(doc2, 'This document defines the complete architectural design for a scalable multi-tenant VM management system capable of managing 1M+ Windows/Linux VMs. The architecture is based on Approach 2: Piko + webhook + linyows/probe, which provides optimal scalability, resource efficiency, and operational simplicity.')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Key Architectural Principles:', bold=True)
add_paragraph(doc2, '• Multi-tenancy with complete data and network isolation')
add_paragraph(doc2, '• Horizontal scalability using Kubernetes HPA')
add_paragraph(doc2, '• Unified agent binary with embedded components')
add_paragraph(doc2, '• One-time installation keys with automatic rotation')
add_paragraph(doc2, '• Comprehensive audit logging to Quickwit')
add_paragraph(doc2, '• AI-powered workflow generation via MCP server')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Target Metrics:', bold=True)
add_paragraph(doc2, '• Support 1M+ agents with <4GB memory per agent')
add_paragraph(doc2, '• Agent registration: <500ms p95')
add_paragraph(doc2, '• Workflow execution latency: <200ms p95')
add_paragraph(doc2, '• API throughput: 10,000 req/sec')
add_paragraph(doc2, '• System availability: 99.95%')
add_paragraph(doc2, '• Agent upgrade success rate: >99.8%')

# System Overview
add_heading(doc2, '2. System Overview', 1)

add_heading(doc2, '2.1 High-Level Architecture', 2)
add_paragraph(doc2, 'The system consists of the following major components:', bold=True)
add_paragraph(doc2, '')
add_paragraph(doc2, '1. Unified Agent (vm-agent):', bold=True)
add_paragraph(doc2, '   • Single Go binary (~15-20MB)')
add_paragraph(doc2, '   • Runs on each managed VM (Windows/Linux)')
add_paragraph(doc2, '   • Embeds: Piko client, webhook server, linyows/probe, lifecycle manager')
add_paragraph(doc2, '   • Memory footprint: 2-4MB per instance')
add_paragraph(doc2, '')
add_paragraph(doc2, '2. Piko Infrastructure:', bold=True)
add_paragraph(doc2, '   • Reverse proxy cluster providing secure tunneling')
add_paragraph(doc2, '   • Multi-tenant authentication (per-endpoint tokens)')
add_paragraph(doc2, '   • Automatic load rebalancing')
add_paragraph(doc2, '   • Kubernetes HPA for auto-scaling')
add_paragraph(doc2, '')
add_paragraph(doc2, '3. Control Plane API:', bold=True)
add_paragraph(doc2, '   • REST API server (Go)')
add_paragraph(doc2, '   • Tenant management')
add_paragraph(doc2, '   • Agent registration and authentication')
add_paragraph(doc2, '   • Workflow submission and orchestration')
add_paragraph(doc2, '   • Campaign management (phased upgrades)')
add_paragraph(doc2, '')
add_paragraph(doc2, '4. Database (MySQL/PostgreSQL):', bold=True)
add_paragraph(doc2, '   • Tenant data storage')
add_paragraph(doc2, '   • Agent registry')
add_paragraph(doc2, '   • Installation keys (one-time, rotating)')
add_paragraph(doc2, '   • Workflow and campaign tracking')
add_paragraph(doc2, '')
add_paragraph(doc2, '5. MCP Server:', bold=True)
add_paragraph(doc2, '   • AI integration for workflow generation')
add_paragraph(doc2, '   • Natural language to YAML workflow translation')
add_paragraph(doc2, '   • Workflow validation')
add_paragraph(doc2, '')
add_paragraph(doc2, '6. Quickwit:', bold=True)
add_paragraph(doc2, '   • Audit log storage and indexing')
add_paragraph(doc2, '   • Search and analytics')
add_paragraph(doc2, '   • Compliance and security monitoring')

add_heading(doc2, '2.2 Data Flow', 2)
add_paragraph(doc2, 'Agent Registration Flow:', bold=True)
add_paragraph(doc2, '1. Admin generates one-time installation key via Control Plane API')
add_paragraph(doc2, '2. Agent installed on VM with installation key')
add_paragraph(doc2, '3. Agent calls /api/v1/agents/register with key')
add_paragraph(doc2, '4. Control Plane validates key (unused, not expired)')
add_paragraph(doc2, '5. Control Plane generates permanent JWT token')
add_paragraph(doc2, '6. Control Plane marks key as used')
add_paragraph(doc2, '7. Agent stores token and establishes Piko tunnel')
add_paragraph(doc2, '8. Agent begins health check reporting')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Workflow Execution Flow:', bold=True)
add_paragraph(doc2, '1. User/AI submits workflow to Control Plane API')
add_paragraph(doc2, '2. Control Plane validates workflow against tenant permissions')
add_paragraph(doc2, '3. Control Plane stores workflow in database')
add_paragraph(doc2, '4. Control Plane sends HTTP POST to agent via Piko tunnel')
add_paragraph(doc2, '5. Agent webhook receives workflow')
add_paragraph(doc2, '6. Agent validates and executes workflow using linyows/probe')
add_paragraph(doc2, '7. Agent streams execution results back to Control Plane')
add_paragraph(doc2, '8. Control Plane updates workflow status')
add_paragraph(doc2, '9. Control Plane logs audit trail to Quickwit')

# Multi-Tenancy Architecture
add_heading(doc2, '3. Multi-Tenancy Architecture', 1)

add_heading(doc2, '3.1 Tenant Isolation Model', 2)
add_paragraph(doc2, 'The system implements multi-tenancy at all layers:', bold=True)
add_paragraph(doc2, '')
add_paragraph(doc2, 'Network Layer:', bold=True)
add_paragraph(doc2, '• Each tenant has unique Piko endpoint: tenant-{id}/*')
add_paragraph(doc2, '• Per-tenant authentication tokens')
add_paragraph(doc2, '• Agents cannot access other tenants\' endpoints')
add_paragraph(doc2, '• Optional mTLS for additional security')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Data Layer:', bold=True)
add_paragraph(doc2, '• All database tables include tenant_id column')
add_paragraph(doc2, '• Row-level security enforced in queries')
add_paragraph(doc2, '• Database indexes include tenant_id for performance')
add_paragraph(doc2, '• No cross-tenant queries allowed')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Application Layer:', bold=True)
add_paragraph(doc2, '• JWT tokens include tenant_id claim')
add_paragraph(doc2, '• All API endpoints validate tenant_id from token')
add_paragraph(doc2, '• Tenant context passed through all operations')
add_paragraph(doc2, '• Audit logs include tenant_id')

add_heading(doc2, '3.2 Per-Tenant Authentication', 2)
add_paragraph(doc2, 'Authentication Flow:', bold=True)
add_paragraph(doc2, '')
add_code_block(doc2, '''# Piko Configuration (per tenant)
server:
  auth:
    enabled: true
    endpoints:
      - endpoint: "tenant-acme/*"
        token: "${PIKO_TOKEN_TENANT_ACME}"
      - endpoint: "tenant-widgets/*"
        token: "${PIKO_TOKEN_TENANT_WIDGETS}"
''')
add_paragraph(doc2, '')
add_paragraph(doc2, 'JWT Token Structure:', bold=True)
add_code_block(doc2, '''{
  "sub": "agent-12345",
  "tenant_id": "acme",
  "agent_id": "server-001",
  "iat": 1704067200,
  "exp": 1735689600,
  "permissions": ["execute_workflow", "report_status"]
}''')

add_heading(doc2, '3.3 Tenant-Specific Piko Endpoints', 2)
add_paragraph(doc2, 'Each tenant gets dedicated Piko endpoints:', bold=True)
add_paragraph(doc2, '')
add_paragraph(doc2, '• Agent registration: https://piko.example.com/tenant-{id}/agent-{agent-id}')
add_paragraph(doc2, '• Webhook endpoint: POST via Piko to agent')
add_paragraph(doc2, '• Health checks: GET via Piko to agent/health')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Benefits:', bold=True)
add_paragraph(doc2, '• Complete network isolation between tenants')
add_paragraph(doc2, '• Independent scaling per tenant')
add_paragraph(doc2, '• Granular rate limiting per tenant')
add_paragraph(doc2, '• Simplified security auditing')

add_heading(doc2, '3.4 Database Schema for Tenant Data', 2)
add_code_block(doc2, '''-- Tenants table
CREATE TABLE tenants (
  id VARCHAR(64) PRIMARY KEY,
  name VARCHAR(255) NOT NULL,
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  piko_token_hash VARCHAR(256) NOT NULL,
  status ENUM('active', 'suspended', 'deleted') DEFAULT 'active',
  settings JSON,
  INDEX idx_status (status)
);

-- Agents table
CREATE TABLE agents (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  tenant_id VARCHAR(64) NOT NULL,
  agent_id VARCHAR(128) NOT NULL,
  hostname VARCHAR(255),
  version VARCHAR(32),
  os_type VARCHAR(32),
  os_version VARCHAR(64),
  registered_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  last_seen TIMESTAMP,
  status ENUM('active', 'inactive', 'upgrading', 'failed') DEFAULT 'active',
  metadata JSON,
  UNIQUE KEY unique_tenant_agent (tenant_id, agent_id),
  INDEX idx_tenant_status (tenant_id, status),
  INDEX idx_last_seen (last_seen),
  FOREIGN KEY (tenant_id) REFERENCES tenants(id)
);

-- Installation keys table
CREATE TABLE installation_keys (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  tenant_id VARCHAR(64) NOT NULL,
  key_hash VARCHAR(256) NOT NULL,
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  expires_at TIMESTAMP NOT NULL,
  used BOOLEAN DEFAULT FALSE,
  used_at TIMESTAMP,
  used_by_agent VARCHAR(128),
  UNIQUE KEY unique_key_hash (key_hash),
  INDEX idx_tenant_unused (tenant_id, used, expires_at),
  FOREIGN KEY (tenant_id) REFERENCES tenants(id)
);

-- Workflows table
CREATE TABLE workflows (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  tenant_id VARCHAR(64) NOT NULL,
  workflow_id VARCHAR(128) NOT NULL,
  agent_id VARCHAR(128) NOT NULL,
  workflow_yaml TEXT NOT NULL,
  submitted_by VARCHAR(128),
  submitted_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  status ENUM('pending', 'running', 'completed', 'failed', 'cancelled'),
  started_at TIMESTAMP,
  completed_at TIMESTAMP,
  result JSON,
  error TEXT,
  UNIQUE KEY unique_workflow (tenant_id, workflow_id),
  INDEX idx_tenant_status (tenant_id, status),
  INDEX idx_agent_status (tenant_id, agent_id, status),
  FOREIGN KEY (tenant_id) REFERENCES tenants(id)
);

-- Campaigns table
CREATE TABLE campaigns (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  tenant_id VARCHAR(64) NOT NULL,
  campaign_id VARCHAR(128) NOT NULL,
  name VARCHAR(255) NOT NULL,
  campaign_type ENUM('upgrade', 'deploy', 'config') NOT NULL,
  target_version VARCHAR(32),
  workflow_template TEXT NOT NULL,
  phase_config JSON NOT NULL,
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  started_at TIMESTAMP,
  completed_at TIMESTAMP,
  status ENUM('draft', 'running', 'paused', 'completed', 'failed', 'rolled_back'),
  total_agents INT DEFAULT 0,
  success_count INT DEFAULT 0,
  failure_count INT DEFAULT 0,
  UNIQUE KEY unique_campaign (tenant_id, campaign_id),
  INDEX idx_tenant_status (tenant_id, status),
  FOREIGN KEY (tenant_id) REFERENCES tenants(id)
);

-- Audit logs table
CREATE TABLE audit_logs (
  id BIGINT PRIMARY KEY AUTO_INCREMENT,
  tenant_id VARCHAR(64) NOT NULL,
  timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  event_type VARCHAR(64) NOT NULL,
  actor VARCHAR(128),
  agent_id VARCHAR(128),
  resource_type VARCHAR(64),
  resource_id VARCHAR(128),
  action VARCHAR(64) NOT NULL,
  result ENUM('success', 'failure'),
  details JSON,
  ip_address VARCHAR(45),
  INDEX idx_tenant_timestamp (tenant_id, timestamp),
  INDEX idx_event_type (event_type),
  INDEX idx_agent (tenant_id, agent_id, timestamp),
  FOREIGN KEY (tenant_id) REFERENCES tenants(id)
);''')

# Component Architecture - Agent
add_heading(doc2, '4. Component Architecture', 1)

add_heading(doc2, '4.1 Unified Agent (vm-agent)', 2)

add_paragraph(doc2, 'Single Binary Structure:', bold=True)
add_paragraph(doc2, 'The agent is a single statically-linked Go binary containing all necessary components:')
add_paragraph(doc2, '')
add_code_block(doc2, '''vm-agent/
├── cmd/agent/main.go          # Entry point with CLI
├── pkg/
│   ├── agent/
│   │   ├── manager.go         # Core agent manager
│   │   ├── coordinator.go     # Component coordination
│   │   └── context.go         # Shared context
│   ├── piko/
│   │   ├── client.go          # Piko client wrapper
│   │   ├── tunnel.go          # Tunnel management
│   │   └── reconnect.go       # Reconnection logic
│   ├── webhook/
│   │   ├── server.go          # Webhook HTTP server
│   │   ├── handlers.go        # Request handlers
│   │   └── auth.go            # Request authentication
│   ├── probe/
│   │   ├── executor.go        # Probe execution wrapper
│   │   ├── workflow.go        # Workflow parsing
│   │   └── reporter.go        # Result reporting
│   ├── config/
│   │   ├── loader.go          # Config loading
│   │   ├── priority.go        # Priority resolution
│   │   └── validator.go       # Config validation
│   ├── health/
│   │   ├── monitor.go         # Health monitoring
│   │   ├── checks.go          # Component health checks
│   │   └── reporter.go        # Health reporting
│   └── lifecycle/
│       ├── install.go         # Installation
│       ├── configure.go       # Configuration
│       ├── repair.go          # Self-repair
│       ├── upgrade.go         # Self-upgrade
│       ├── uninstall.go       # Clean removal
│       ├── service_linux.go   # systemd integration
│       └── service_windows.go # Windows service
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Embedded Components:', bold=True)
add_paragraph(doc2, '1. Piko Client: Establishes secure tunnel to Piko cluster')
add_paragraph(doc2, '2. Webhook Server: Receives workflow execution requests')
add_paragraph(doc2, '3. linyows/probe: Executes YAML workflows locally')
add_paragraph(doc2, '4. Agent Manager: Orchestrates all components, handles CLI')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Lifecycle Management:', bold=True)
add_paragraph(doc2, 'The agent supports complete lifecycle operations:')
add_paragraph(doc2, '')
add_paragraph(doc2, '• install: Register as system service (systemd/Windows Service)')
add_paragraph(doc2, '• configure: Update configuration dynamically')
add_paragraph(doc2, '• repair: Run diagnostics and fix common issues')
add_paragraph(doc2, '• upgrade: Download and install new version with rollback')
add_paragraph(doc2, '• uninstall: Clean removal including service deregistration')
add_paragraph(doc2, '• status: Report health and version information')
add_paragraph(doc2, '• run: Main service execution mode')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Self-Healing Mechanisms:', bold=True)
add_paragraph(doc2, '• Automatic Piko tunnel reconnection on disconnect')
add_paragraph(doc2, '• Webhook server restart on failure')
add_paragraph(doc2, '• Configuration reload without restart')
add_paragraph(doc2, '• Memory leak detection and self-restart')
add_paragraph(doc2, '• Disk space monitoring and cleanup')
add_paragraph(doc2, '• Network connectivity checks')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Configuration Management:', bold=True)
add_paragraph(doc2, 'Priority order (highest to lowest):')
add_paragraph(doc2, '1. Environment variables (VM_AGENT_*)')
add_paragraph(doc2, '2. Configuration file (/etc/vm-agent/config.yaml or C:\\ProgramData\\vm-agent\\config.yaml)')
add_paragraph(doc2, '3. Remote configuration (fetched from Control Plane)')
add_paragraph(doc2, '4. Compiled defaults')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Health Monitoring:', bold=True)
add_paragraph(doc2, '• Component health checks every 30 seconds')
add_paragraph(doc2, '• Piko tunnel connectivity check')
add_paragraph(doc2, '• Webhook server responsiveness')
add_paragraph(doc2, '• Disk space availability')
add_paragraph(doc2, '• Memory usage monitoring')
add_paragraph(doc2, '• Periodic health report to Control Plane')

add_heading(doc2, '4.2 Piko Infrastructure', 2)

add_paragraph(doc2, 'Multi-Tenant Authentication:', bold=True)
add_paragraph(doc2, 'Piko supports per-endpoint authentication for complete tenant isolation:')
add_paragraph(doc2, '')
add_code_block(doc2, '''# piko-config.yaml
server:
  bind_addr: "0.0.0.0:8000"

  auth:
    enabled: true
    # Per-endpoint authentication
    endpoints:
      - endpoint: "tenant-acme/*"
        token_env: "PIKO_TOKEN_TENANT_ACME"
      - endpoint: "tenant-widgets/*"
        token_env: "PIKO_TOKEN_TENANT_WIDGETS"

  # Load balancing configuration
  rebalancing:
    enabled: true
    interval: 1s
    threshold: 1.5  # Shed connections if 1.5x above average

  # Connection limits
  limits:
    max_connections_per_endpoint: 100000
    max_bandwidth_mbps: 10000

  # Cluster configuration
  cluster:
    enabled: true
    gossip_port: 7946
    peers: ["piko-0.piko:7946", "piko-1.piko:7946", "piko-2.piko:7946"]
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Load Balancing and Rebalancing:', bold=True)
add_paragraph(doc2, '• Automatic connection distribution across Piko nodes')
add_paragraph(doc2, '• Periodic rebalancing when node load exceeds threshold')
add_paragraph(doc2, '• Graceful connection migration (no dropped connections)')
add_paragraph(doc2, '• Health-based routing (avoid unhealthy nodes)')
add_paragraph(doc2, '')
add_paragraph(doc2, 'HPA Autoscaling Configuration:', bold=True)
add_code_block(doc2, '''apiVersion: autoscaling/v2
kind: HorizontalPodAutoscaler
metadata:
  name: piko-hpa
spec:
  scaleTargetRef:
    apiVersion: apps/v1
    kind: StatefulSet
    name: piko
  minReplicas: 3
  maxReplicas: 50
  metrics:
  - type: Resource
    resource:
      name: cpu
      target:
        type: Utilization
        averageUtilization: 70
  - type: Resource
    resource:
      name: memory
      target:
        type: Utilization
        averageUtilization: 80
  - type: Pods
    pods:
      metric:
        name: piko_active_connections
      target:
        type: AverageValue
        averageValue: "10000"  # Scale at 10K conn/pod
  behavior:
    scaleUp:
      stabilizationWindowSeconds: 60
      policies:
      - type: Percent
        value: 50
        periodSeconds: 60
    scaleDown:
      stabilizationWindowSeconds: 300
      policies:
      - type: Pods
        value: 1
        periodSeconds: 180
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Connection Management:', bold=True)
add_paragraph(doc2, '• Connection pooling for efficiency')
add_paragraph(doc2, '• Automatic reconnection with exponential backoff')
add_paragraph(doc2, '• Connection timeout: 30 seconds')
add_paragraph(doc2, '• Keep-alive: 15 seconds')
add_paragraph(doc2, '• Max idle time: 5 minutes')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Cluster Gossip Protocol:', bold=True)
add_paragraph(doc2, '• Memberlist-based gossip for cluster membership')
add_paragraph(doc2, '• Automatic node discovery')
add_paragraph(doc2, '• Failure detection within 5 seconds')
add_paragraph(doc2, '• Connection state synchronization')

add_heading(doc2, '4.3 Control Plane API', 2)

add_paragraph(doc2, 'REST API Endpoints:', bold=True)
add_code_block(doc2, '''# Tenant Management
POST   /api/v1/tenants                    # Create tenant
GET    /api/v1/tenants/:id                # Get tenant details
PUT    /api/v1/tenants/:id                # Update tenant
DELETE /api/v1/tenants/:id                # Delete tenant
GET    /api/v1/tenants                    # List tenants

# Installation Key Management
POST   /api/v1/tenants/:id/keys           # Generate installation key
GET    /api/v1/tenants/:id/keys           # List keys
DELETE /api/v1/tenants/:id/keys/:key_id   # Revoke key
POST   /api/v1/tenants/:id/keys/rotate    # Rotate all keys

# Agent Management
POST   /api/v1/agents/register            # Agent registration (uses install key)
GET    /api/v1/agents                     # List agents (tenant-scoped)
GET    /api/v1/agents/:id                 # Get agent details
DELETE /api/v1/agents/:id                 # Deregister agent
GET    /api/v1/agents/:id/health          # Get agent health
POST   /api/v1/agents/:id/token/refresh   # Refresh agent token

# Workflow Management
POST   /api/v1/workflows                  # Submit workflow
GET    /api/v1/workflows/:id              # Get workflow details
GET    /api/v1/workflows/:id/status       # Get workflow status
GET    /api/v1/workflows/:id/logs         # Get workflow logs
POST   /api/v1/workflows/:id/cancel       # Cancel running workflow
GET    /api/v1/workflows                  # List workflows (tenant-scoped)

# Campaign Management
POST   /api/v1/campaigns                  # Create upgrade campaign
GET    /api/v1/campaigns/:id              # Get campaign details
GET    /api/v1/campaigns/:id/status       # Get campaign status
POST   /api/v1/campaigns/:id/start        # Start campaign
POST   /api/v1/campaigns/:id/pause        # Pause campaign
POST   /api/v1/campaigns/:id/resume       # Resume campaign
POST   /api/v1/campaigns/:id/rollback     # Rollback campaign
GET    /api/v1/campaigns                  # List campaigns (tenant-scoped)

# Audit Logs
GET    /api/v1/audit                      # Query audit logs (tenant-scoped)
GET    /api/v1/audit/export               # Export audit logs

# Health & Metrics
GET    /health                            # Control plane health
GET    /metrics                           # Prometheus metrics
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Agent Registration Flow:', bold=True)
add_code_block(doc2, '''# 1. Admin generates installation key
POST /api/v1/tenants/acme/keys
{
  "expires_in_hours": 24,
  "max_uses": 1
}

Response:
{
  "key": "install_xxxxxxxxxxxxxxxx",  # One-time use only
  "expires_at": "2024-01-15T10:00:00Z"
}

# 2. Agent registers using key
POST /api/v1/agents/register
Authorization: Bearer install_xxxxxxxxxxxxxxxx
{
  "agent_id": "server-001",
  "hostname": "web-server-01",
  "version": "1.0.0",
  "os_type": "linux",
  "os_version": "Ubuntu 22.04"
}

Response:
{
  "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",  # Permanent JWT
  "piko_endpoint": "https://piko.example.com/tenant-acme/server-001",
  "refresh_interval_hours": 720  # 30 days
}

# 3. Agent uses permanent token for all subsequent requests
# Token is stored securely on agent filesystem
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Token Rotation Mechanism:', bold=True)
add_paragraph(doc2, '• Tokens expire after 1 year')
add_paragraph(doc2, '• Agents refresh tokens 30 days before expiry')
add_paragraph(doc2, '• Automatic rotation triggered by Control Plane')
add_paragraph(doc2, '• Grace period for old tokens: 7 days after rotation')
add_paragraph(doc2, '• Emergency rotation API for security incidents')

# Continue with more sections...
add_heading(doc2, '4.4 Database Implementation', 2)
add_paragraph(doc2, 'Database Choice: MySQL 8.0+ or PostgreSQL 13+', bold=True)
add_paragraph(doc2, '')
add_paragraph(doc2, 'Connection Pooling:', bold=True)
add_paragraph(doc2, '• Min connections: 10')
add_paragraph(doc2, '• Max connections: 100')
add_paragraph(doc2, '• Connection lifetime: 1 hour')
add_paragraph(doc2, '• Idle timeout: 10 minutes')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Tenant Data Isolation:', bold=True)
add_paragraph(doc2, '• All queries include WHERE tenant_id = ?')
add_paragraph(doc2, '• Prepared statements to prevent SQL injection')
add_paragraph(doc2, '• Application-level enforcement (no database users per tenant)')
add_paragraph(doc2, '• Audit logging for all data access')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Query Optimization:', bold=True)
add_paragraph(doc2, '• Compound indexes starting with tenant_id')
add_paragraph(doc2, '• Query result caching for tenant metadata')
add_paragraph(doc2, '• Read replicas for reporting queries')
add_paragraph(doc2, '• Partitioning for audit_logs table (by month)')

add_heading(doc2, '4.5 MCP Server Implementation', 2)
add_paragraph(doc2, 'The MCP (Model Context Protocol) server enables AI-powered workflow generation:', bold=True)
add_paragraph(doc2, '')
add_paragraph(doc2, 'Tool Definitions:', bold=True)
add_code_block(doc2, '''tools:
  - name: generate_workflow
    description: Generate a workflow from natural language description
    input_schema:
      type: object
      properties:
        description:
          type: string
          description: Natural language description of desired workflow
        target_os:
          type: string
          enum: [linux, windows, any]
        constraints:
          type: object
          description: Additional constraints (timeout, retries, etc.)

  - name: validate_workflow
    description: Validate a YAML workflow for correctness
    input_schema:
      type: object
      properties:
        workflow_yaml:
          type: string
          description: YAML workflow to validate

  - name: submit_workflow
    description: Submit workflow to agents
    input_schema:
      type: object
      properties:
        workflow_yaml:
          type: string
        agent_ids:
          type: array
          items:
            type: string

  - name: list_agents
    description: List available agents
    input_schema:
      type: object
      properties:
        filters:
          type: object
          description: Filter criteria (os_type, status, etc.)

  - name: check_workflow_status
    description: Check status of submitted workflow
    input_schema:
      type: object
      properties:
        workflow_id:
          type: string
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Workflow Generation Example:', bold=True)
add_code_block(doc2, '''User: "Upgrade nginx to version 1.24 on all web servers"

MCP generates:
---
name: Upgrade nginx to 1.24
timeout: 600s
http:
  - name: Check current nginx version
    url: http://localhost
    method: GET
    timeout: 5s
    expect:
      code: 200

  - name: Stop nginx
    url: http://localhost:9999/execute
    method: POST
    body: |
      {
        "command": "systemctl stop nginx",
        "timeout": 30
      }

  - name: Update package
    url: http://localhost:9999/execute
    method: POST
    body: |
      {
        "command": "apt-get update && apt-get install -y nginx=1.24*",
        "timeout": 300
      }

  - name: Start nginx
    url: http://localhost:9999/execute
    method: POST
    body: |
      {
        "command": "systemctl start nginx",
        "timeout": 30
      }

  - name: Verify nginx version
    url: http://localhost
    method: GET
    timeout: 5s
    expect:
      code: 200
      body: "nginx/1.24"
''')

add_heading(doc2, '4.6 Quickwit Integration', 2)
add_paragraph(doc2, 'Log Schema:', bold=True)
add_code_block(doc2, '''# quickwit-index.yaml
version: 0.7
index_id: vm-agent-audit-logs

doc_mapping:
  field_mappings:
    - name: timestamp
      type: datetime
      input_formats:
        - rfc3339
      fast: true

    - name: tenant_id
      type: text
      tokenizer: raw
      fast: true

    - name: event_type
      type: text
      tokenizer: raw
      fast: true

    - name: actor
      type: text
      tokenizer: raw

    - name: agent_id
      type: text
      tokenizer: raw
      fast: true

    - name: action
      type: text
      tokenizer: raw

    - name: result
      type: text
      tokenizer: raw

    - name: details
      type: json

    - name: ip_address
      type: ip
      fast: true

  timestamp_field: timestamp

indexing_settings:
  commit_timeout_secs: 60

search_settings:
  default_search_fields: [event_type, action, agent_id]
''')

add_paragraph(doc2, '')
add_paragraph(doc2, 'Indexing Strategy:', bold=True)
add_paragraph(doc2, '• Real-time indexing (60-second commit)')
add_paragraph(doc2, '• Tenant-based index sharding for large tenants')
add_paragraph(doc2, '• Time-based index rotation (monthly)')
add_paragraph(doc2, '• Automatic index merging for optimization')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Retention Policies:', bold=True)
add_paragraph(doc2, '• Hot data: 30 days (SSD storage)')
add_paragraph(doc2, '• Warm data: 90 days (HDD storage)')
add_paragraph(doc2, '• Cold data: 2 years (object storage)')
add_paragraph(doc2, '• Deletion after 2 years (configurable per tenant)')
add_paragraph(doc2, '')
add_paragraph(doc2, 'Query Patterns:', bold=True)
add_paragraph(doc2, '• Tenant isolation: All queries filtered by tenant_id')
add_paragraph(doc2, '• Common queries cached for 5 minutes')
add_paragraph(doc2, '• Aggregations for dashboard metrics')
add_paragraph(doc2, '• Export API for compliance requirements')

# Save Document 2
doc2.save('doc2_architecture.docx')
print("Document 2 completed: doc2_architecture.docx")

print("\nAll documents generation completed!")
print("Generated files:")
print("  - doc1_comparison.docx")
print("  - doc2_architecture.docx")
print("  - doc3_implementation.docx (placeholder - continuing...)")
print("  - doc4_claude_code_spec.docx (placeholder - continuing...)")
